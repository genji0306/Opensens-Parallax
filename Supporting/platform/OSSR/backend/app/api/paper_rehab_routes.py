"""
Paper Rehabilitation Routes
Upload paper drafts, run adversarial AI review/revision, track progress.

Endpoints:
  POST /paper-lab/upload              — Upload .docx/.txt/.md, parse, return upload_id
  GET  /paper-lab/<id>/status         — Current status + scores
  POST /paper-lab/<id>/start-review   — Begin review game (async)
  POST /paper-lab/<id>/specialist-review — Run specialist domain review (async)
  GET  /paper-lab/<id>/specialist-review  — Fetch latest specialist review snapshot
  GET  /paper-lab/<id>/rounds         — Fetch all review/revision round data
  GET  /paper-lab/<id>/draft          — Current draft text
  GET  /paper-lab/<id>/stream         — SSE stream for live progress
  GET  /paper-lab/uploads             — List all uploads
"""

import json
import logging
import os
import re
import threading
import time
import uuid
from datetime import datetime
from pathlib import Path

from flask import Blueprint, request, jsonify, Response
from opensens_common.config import Config
from opensens_common.task import TaskManager, TaskStatus

from ..db import get_connection

logger = logging.getLogger(__name__)

paper_rehab_bp = Blueprint("paper_rehab", __name__)

# Upload directory for paper files
UPLOAD_DIR = Path(__file__).parent.parent.parent / "data" / "paper_uploads"

# SSE progress events stored per upload_id
_sse_events = {}  # upload_id -> list of event dicts


# ── Helpers ────────────────────────────────────────────────────────


def _save_upload(upload_id: str, data: dict):
    """Insert or replace a paper_uploads row."""
    conn = get_connection()
    conn.execute(
        """INSERT OR REPLACE INTO paper_uploads
        (upload_id, title, language, detected_field, sections, raw_references,
         full_text, metadata, source_filename, status, review_config,
         review_rounds, source_audit, reviewers, authors, current_draft,
         score_progression, created_at, updated_at, error)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            upload_id,
            data.get("title", ""),
            data.get("language", "en"),
            data.get("detected_field", ""),
            json.dumps(data.get("sections", [])),
            json.dumps(data.get("raw_references", [])),
            data.get("full_text", ""),
            json.dumps(data.get("metadata", {})),
            data.get("source_filename", ""),
            data.get("status", "uploaded"),
            json.dumps(data.get("review_config", {})),
            json.dumps(data.get("review_rounds", [])),
            json.dumps(data.get("source_audit", {})),
            json.dumps(data.get("reviewers", [])),
            json.dumps(data.get("authors", [])),
            data.get("current_draft", ""),
            json.dumps(data.get("score_progression", [])),
            data.get("created_at", datetime.now().isoformat()),
            datetime.now().isoformat(),
            data.get("error"),
        ),
    )
    conn.commit()


def _load_upload(upload_id: str) -> dict | None:
    """Load a paper_uploads row as dict."""
    conn = get_connection()
    row = conn.execute("SELECT * FROM paper_uploads WHERE upload_id = ?", (upload_id,)).fetchone()
    if not row:
        return None
    return {
        "upload_id": row["upload_id"],
        "title": row["title"],
        "language": row["language"],
        "detected_field": row["detected_field"],
        "sections": json.loads(row["sections"]) if row["sections"] else [],
        "raw_references": json.loads(row["raw_references"]) if row["raw_references"] else [],
        "full_text": row["full_text"],
        "metadata": json.loads(row["metadata"]) if row["metadata"] else {},
        "source_filename": row["source_filename"],
        "status": row["status"],
        "review_config": json.loads(row["review_config"]) if row["review_config"] else {},
        "review_rounds": json.loads(row["review_rounds"]) if row["review_rounds"] else [],
        "source_audit": json.loads(row["source_audit"]) if row["source_audit"] else {},
        "reviewers": json.loads(row["reviewers"]) if row["reviewers"] else [],
        "authors": json.loads(row["authors"]) if row["authors"] else [],
        "current_draft": row["current_draft"],
        "score_progression": json.loads(row["score_progression"]) if row["score_progression"] else [],
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
        "error": row["error"],
    }


def _current_review_session_id(upload: dict) -> str:
    meta = upload.get("metadata") or {}
    return meta.get("current_review_session_id", "")


def _update_upload_metadata(upload_id: str, mutate_fn):
    upload = _load_upload(upload_id)
    if not upload:
        return None
    meta = dict(upload.get("metadata") or {})
    mutate_fn(meta)
    conn = get_connection()
    conn.execute(
        "UPDATE paper_uploads SET metadata = ?, updated_at = ? WHERE upload_id = ?",
        (json.dumps(meta), datetime.now().isoformat(), upload_id),
    )
    conn.commit()
    upload["metadata"] = meta
    return upload


def _build_draft_from_sections(sections: list[dict]) -> str:
    blocks = []
    for section in sections:
        name = (section.get("name") or "").strip()
        content = (section.get("content") or section.get("text") or "").strip()
        if name:
            blocks.append(name)
        if content:
            blocks.append(content)
    return "\n\n".join(block for block in blocks if block).strip()


def _apply_section_revision_to_upload(upload: dict, section_name: str, revised_text: str, action: str, refinement_id: str) -> dict:
    sections = list(upload.get("sections") or [])
    previous_draft = upload.get("current_draft", "")
    updated = False
    previous_section_text = ""
    for section in sections:
        if (section.get("name") or "").lower() == section_name.lower():
            previous_section_text = section.get("content") or section.get("text") or ""
            if "content" in section:
                section["content"] = revised_text
            else:
                section["text"] = revised_text
            updated = True
            break

    if not updated:
        sections.append({"name": section_name, "content": revised_text})

    next_draft = _build_draft_from_sections(sections) or revised_text
    metadata = dict(upload.get("metadata") or {})
    history = list(metadata.get("applied_section_refinements") or [])
    history.append({
        "refinement_id": refinement_id,
        "action": action,
        "section": section_name,
        "applied_at": datetime.now().isoformat(),
        "section_before_text": previous_section_text,
        "section_after_text": revised_text,
        "draft_before_excerpt": previous_draft[:280],
        "draft_after_excerpt": next_draft[:280],
    })
    metadata["applied_section_refinements"] = history[-20:]

    conn = get_connection()
    conn.execute(
        """
        UPDATE paper_uploads
        SET sections = ?, current_draft = ?, metadata = ?, updated_at = ?
        WHERE upload_id = ?
        """,
        (
            json.dumps(sections),
            next_draft,
            json.dumps(metadata),
            datetime.now().isoformat(),
            upload["upload_id"],
        ),
    )
    conn.commit()

    upload["sections"] = sections
    upload["current_draft"] = next_draft
    upload["metadata"] = metadata
    return upload


def _revert_section_refinement(upload: dict, refinement_id: str) -> dict | None:
    metadata = dict(upload.get("metadata") or {})
    history = list(metadata.get("applied_section_refinements") or [])
    target = next((item for item in reversed(history) if item.get("refinement_id") == refinement_id), None)
    if not target:
        return None

    section_name = target.get("section") or ""
    previous_text = target.get("section_before_text", "")
    sections = list(upload.get("sections") or [])
    for section in sections:
        if (section.get("name") or "").lower() == section_name.lower():
            if "content" in section:
                section["content"] = previous_text
            else:
                section["text"] = previous_text
            break

    next_draft = _build_draft_from_sections(sections)
    history.append({
        "refinement_id": f"revert_{uuid.uuid4().hex[:10]}",
        "action": f"revert:{target.get('action', 'unknown')}",
        "section": section_name,
        "applied_at": datetime.now().isoformat(),
        "section_before_text": target.get("section_after_text", ""),
        "section_after_text": previous_text,
        "draft_before_excerpt": (upload.get("current_draft") or "")[:280],
        "draft_after_excerpt": next_draft[:280],
        "reverted_refinement_id": refinement_id,
    })
    metadata["applied_section_refinements"] = history[-20:]
    metadata["last_reverted_refinement_id"] = refinement_id

    section_history = []
    for item in (metadata.get("section_refinement_history") or []):
        if item.get("refinement_id") == refinement_id:
            section_history.append({
                **item,
                "applied": False,
                "reverted_at": datetime.now().isoformat(),
            })
        else:
            section_history.append(item)
    metadata["section_refinement_history"] = section_history

    conn = get_connection()
    conn.execute(
        """
        UPDATE paper_uploads
        SET sections = ?, current_draft = ?, metadata = ?, updated_at = ?
        WHERE upload_id = ?
        """,
        (
            json.dumps(sections),
            next_draft,
            json.dumps(metadata),
            datetime.now().isoformat(),
            upload["upload_id"],
        ),
    )
    conn.commit()

    upload["sections"] = sections
    upload["current_draft"] = next_draft
    upload["metadata"] = metadata
    return upload


def _artifact_export_bundle(artifact: dict) -> dict:
    payload = artifact.get("payload") or {}
    rendering = payload.get("rendering") or {}
    export_formats = list(dict.fromkeys(payload.get("export_formats") or ["json"]))
    base_name = (artifact.get("title", "artifact").strip() or "artifact").replace(" ", "_").lower()
    artifact_type = artifact.get("type")
    spec = rendering.get("spec")
    files = [
        {
            "filename": f"{base_name}.json",
            "format": "json",
            "content": artifact,
        }
    ]

    if rendering.get("spec") is not None:
        files.append({
            "filename": f"{base_name}.spec.json",
            "format": "json",
            "content": {
                "engine": rendering.get("engine"),
                "spec": rendering.get("spec"),
            },
        })

    if artifact_type == "graphical_abstract" and isinstance(spec, str):
        files.append({
            "filename": f"{base_name}.html",
            "format": "html",
            "content": spec,
        })

    if artifact_type in {"chart", "diagram"} and spec is not None:
        files.append({
            "filename": f"{base_name}.svg",
            "format": "svg",
            "content": (
                f"<svg xmlns='http://www.w3.org/2000/svg' width='960' height='540'>"
                f"<rect width='100%' height='100%' fill='white'/>"
                f"<text x='40' y='80' font-size='28' fill='#111827'>{artifact.get('title', 'Artifact')}</text>"
                f"<text x='40' y='130' font-size='16' fill='#4b5563'>Engine: {rendering.get('engine', 'unknown')}</text>"
                f"<text x='40' y='170' font-size='14' fill='#6b7280'>Exported placeholder bundle for {artifact_type}</text>"
                f"</svg>"
            ),
        })

    if artifact_type == "slide" and isinstance(payload.get("slides"), list):
        files.append({
            "filename": f"{base_name}.slides.json",
            "format": "json",
            "content": payload.get("slides"),
        })
        html_slides = "".join(
            f"<section><h2>{slide.get('title', 'Slide')}</h2><p>{slide.get('summary', '')}</p></section>"
            for slide in payload.get("slides", [])
        )
        files.append({
            "filename": f"{base_name}.html",
            "format": "html",
            "content": f"<article class='slides'>{html_slides}</article>",
        })

    if artifact_type == "poster_panel" and isinstance(payload.get("panels"), list):
        files.append({
            "filename": f"{base_name}.poster.json",
            "format": "json",
            "content": payload.get("panels"),
        })
        html_panels = "".join(
            f"<section><h2>{panel.get('name', 'Panel')}</h2><p>{' | '.join(panel.get('content', [])) if isinstance(panel.get('content'), list) else panel.get('content', '')}</p></section>"
            for panel in payload.get("panels", [])
        )
        files.append({
            "filename": f"{base_name}.html",
            "format": "html",
            "content": f"<article class='poster'>{html_panels}</article>",
        })

    return {
        "bundle_name": base_name,
        "formats": export_formats,
        "files": files,
        "provenance": artifact.get("provenance") or {},
    }


def _emit_sse(upload_id: str, event_type: str, data: dict, session_id: str = ""):
    """Push an SSE event for a given upload_id."""
    if upload_id not in _sse_events:
        _sse_events[upload_id] = []
    _sse_events[upload_id].append({
        "type": event_type,
        "data": data,
        "session_id": session_id,
        "timestamp": datetime.now().isoformat(),
    })


def _summarize_upload(upload: dict) -> dict:
    """Normalize a paper upload for list/status responses."""
    scores = list(upload.get("score_progression", []) or [])
    rounds = len(upload.get("review_rounds", []) or []) or len(scores)
    field = upload.get("detected_field", "") or upload.get("field", "") or "general"
    metadata = upload.get("metadata") or {}
    parse_quality = metadata.get("parse_quality", {}) or {}
    document_counts = {
        "sections": len(upload.get("sections") or []),
        "tables": len(metadata.get("tables") or []),
        "figures": len(metadata.get("figures") or []),
        "formulas": len(metadata.get("formulas") or []),
        "references": len(upload.get("raw_references") or []),
    }

    return {
        "upload_id": upload["upload_id"],
        "title": upload.get("title", ""),
        "language": upload.get("language", "en"),
        "field": field,
        "detected_field": field,
        "status": upload.get("status", "uploaded"),
        "source_filename": upload.get("source_filename", ""),
        "review_scores": scores,
        "score_progression": scores,
        "initial_score": scores[0] if scores else None,
        "final_score": scores[-1] if scores else None,
        "round_count": rounds,
        "rounds_completed": rounds,
        "created_at": upload.get("created_at", ""),
        "updated_at": upload.get("updated_at", ""),
        "parser_engine": metadata.get("parser_engine", metadata.get("parser", "unknown")),
        "parser_mode": metadata.get("parser_mode", "unknown"),
        "parse_quality": parse_quality,
        "ocr_used": bool(metadata.get("ocr_used", False)),
        "document_counts": document_counts,
    }


# ── Upload ─────────────────────────────────────────────────────────


@paper_rehab_bp.route("/paper-lab/upload", methods=["POST"])
def upload_paper():
    """
    Upload a paper draft file (.pdf, .doc, .docx, .txt, .md).
    Accepts multipart/form-data with a 'file' field.
    Returns parsed paper metadata + upload_id for subsequent operations.
    """
    if "file" not in request.files:
        return jsonify({"success": False, "error": "No file provided. Use multipart form with 'file' field."}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"success": False, "error": "Empty filename"}), 400

    ext = Path(file.filename).suffix.lower()
    if ext not in (".pdf", ".doc", ".docx", ".txt", ".md", ".markdown"):
        return jsonify({"success": False, "error": f"Unsupported file type: {ext}. Use .pdf, .doc, .docx, .txt, or .md"}), 400

    # Save uploaded file
    upload_id = f"paper_{uuid.uuid4().hex[:10]}"
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    saved_path = UPLOAD_DIR / f"{upload_id}{ext}"
    file.save(str(saved_path))

    # Parse
    try:
        from ..services.ais.paper_parser import PaperParser
        parser = PaperParser()
        parsed = parser.parse(str(saved_path))

        upload_data = {
            "title": parsed.title,
            "language": parsed.language,
            "detected_field": parsed.detected_field,
            "sections": [s.to_dict() for s in parsed.sections],
            "raw_references": parsed.raw_references,
            "full_text": parsed.full_text,
            "metadata": parsed.metadata,
            "source_filename": file.filename,
            "status": "parsed",
            "current_draft": parsed.full_text,
            "created_at": datetime.now().isoformat(),
        }
        _save_upload(upload_id, upload_data)

        return jsonify({
            "success": True,
            "data": {
                "upload_id": upload_id,
                "title": parsed.title,
                "language": parsed.language,
                "detected_field": parsed.detected_field,
                "section_count": len(parsed.sections),
                "sections": [s.name for s in parsed.sections],
                "word_count": parsed.metadata.get("word_count", 0),
                "reference_count": len(parsed.raw_references),
                "parser_engine": parsed.metadata.get("parser_engine", parsed.metadata.get("parser", "unknown")),
                "parser_mode": parsed.metadata.get("parser_mode", "unknown"),
                "parse_quality": parsed.metadata.get("parse_quality", {}),
            },
        })
    except Exception as e:
        logger.exception("Paper parse failed")
        return jsonify({"success": False, "error": str(e)}), 500


# ── Start Review Game ──────────────────────────────────────────────


@paper_rehab_bp.route("/paper-lab/<upload_id>/start-review", methods=["POST"])
def start_review(upload_id):
    """
    Begin the adversarial review game (async).
    Body: { "rounds": 3, "reviewers": 5, "authors": 3, "live": true }
    Returns task_id for polling.
    """
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    if upload["status"] not in ("parsed", "review_complete", "review_failed"):
        return jsonify({"success": False, "error": f"Cannot start review: status is '{upload['status']}'"}), 400

    data = request.get_json() or {}
    rounds = data.get("rounds", 3)
    num_reviewers = data.get("reviewers", 5)
    num_authors = data.get("authors", 3)
    use_live = data.get("live", False)

    config = {
        "rounds": rounds,
        "reviewers": num_reviewers,
        "authors": num_authors,
        "live": use_live,
    }

    session_id = f"review_{uuid.uuid4().hex[:12]}"
    latest_meta = upload.get("metadata") or {}
    latest_meta["current_review_session_id"] = session_id
    _sse_events[upload_id] = []

    # Update status
    conn = get_connection()
    conn.execute(
        "UPDATE paper_uploads SET status = ?, review_config = ?, metadata = ?, updated_at = ? WHERE upload_id = ?",
        ("reviewing", json.dumps(config), json.dumps(latest_meta), datetime.now().isoformat(), upload_id),
    )
    conn.commit()

    # Create async task
    tm = TaskManager()
    task_id = tm.create_task("paper_review", metadata={"upload_id": upload_id})

    # Launch background thread
    thread = threading.Thread(
        target=_run_review_pipeline,
        args=(upload_id, task_id, config, session_id),
        daemon=True,
    )
    thread.start()

    return jsonify({
        "success": True,
        "data": {
            "upload_id": upload_id,
            "task_id": task_id,
            "session_id": session_id,
            "config": config,
        },
    }), 202


def _run_review_pipeline(upload_id: str, task_id: str, config: dict, session_id: str):
    """Background thread: runs the full review/revision pipeline."""
    import sys
    sys.path.insert(0, str(Path(__file__).parent.parent.parent))

    from app.db import init_db
    init_db()

    tm = TaskManager()
    tm.update_task(task_id, status=TaskStatus.PROCESSING, message="Starting review pipeline")

    try:
        upload = _load_upload(upload_id)
        if not upload:
            raise ValueError("Upload not found")

        rounds = config.get("rounds", 3)
        num_reviewers = config.get("reviewers", 5)
        num_authors = config.get("authors", 3)
        use_live = config.get("live", False)

        # Import the CLI pipeline functions
        sys.path.insert(0, str(Path(__file__).parent.parent.parent))
        from cli_test_paper_rehab import (
            REVIEWER_ARCHETYPES, AUTHOR_ARCHETYPES,
            test_source_audit, test_review_round, test_revision_round,
        )
        from app.services.ais.paper_parser import ParsedPaper, PaperParser
        from app.models.ais_models import PaperSection

        # Reconstruct ParsedPaper from stored data
        sections = [PaperSection.from_dict(s) for s in upload["sections"]]
        parsed_paper = ParsedPaper(
            source_path=upload.get("source_filename", ""),
            title=upload["title"],
            language=upload["language"],
            detected_field=upload["detected_field"],
            sections=sections,
            raw_references=upload["raw_references"],
            metadata=upload["metadata"],
            full_text=upload["full_text"],
        )

        reviewers = REVIEWER_ARCHETYPES[:num_reviewers]
        authors = AUTHOR_ARCHETYPES[:num_authors]
        section_names = [s.name for s in sections]
        current_draft = upload["current_draft"] or upload["full_text"]

        _emit_sse(upload_id, "status", {"message": "Generating reviewer and author agents"}, session_id)
        tm.update_task(task_id, progress=5, message="Agents generated")

        # Source audit
        _emit_sse(upload_id, "status", {"message": "Running source quality audit"}, session_id)
        audit_result = test_source_audit(parsed_paper, use_live)
        source_audit = audit_result.get("_source_audit", {})
        tm.update_task(task_id, progress=15, message="Source audit complete")
        _emit_sse(upload_id, "source_audit", {
            "verified": len(source_audit.get("verified", [])),
            "unverified": len(source_audit.get("unverified", [])),
        }, session_id)

        # Save audit + agents
        conn = get_connection()
        conn.execute(
            "UPDATE paper_uploads SET source_audit = ?, reviewers = ?, authors = ?, updated_at = ? WHERE upload_id = ?",
            (json.dumps(source_audit, default=str), json.dumps(reviewers), json.dumps(authors),
             datetime.now().isoformat(), upload_id),
        )
        conn.commit()

        all_consolidations = []
        all_revisions = []
        all_rounds_data = []
        score_progression = []

        for round_num in range(1, rounds + 1):
            progress_base = 15 + (round_num - 1) * (70 // rounds)

            # Review round
            _emit_sse(upload_id, "review_start", {"round": round_num}, session_id)
            tm.update_task(task_id, progress=progress_base, message=f"Review round {round_num}")

            review_result = test_review_round(
                round_num=round_num,
                current_draft=current_draft,
                reviewers=reviewers,
                section_names=section_names,
                source_audit=source_audit,
                use_live=use_live,
                prev_reviews=all_consolidations,
                prev_revisions=all_revisions,
            )
            consolidated = review_result.get("_consolidated", {})
            all_consolidations.append(consolidated)

            avg_score = consolidated.get("avg_overall_score", 0)
            decision = consolidated.get("final_decision", "unknown")
            score_progression.append(avg_score)

            _emit_sse(upload_id, "review_complete", {
                "round": round_num,
                "avg_score": avg_score,
                "decision": decision,
            }, session_id)

            # Revision round
            _emit_sse(upload_id, "revision_start", {"round": round_num}, session_id)
            tm.update_task(task_id, progress=progress_base + (35 // rounds),
                           message=f"Revision round {round_num}")

            revision_result = test_revision_round(
                round_num=round_num,
                current_draft=current_draft,
                consolidated=consolidated,
                paper_field=parsed_paper.detected_field,
                paper_language=parsed_paper.language,
                use_live=use_live,
            )
            revision = revision_result.get("_revision", {})
            all_revisions.append(revision)

            if revision.get("revised_draft"):
                current_draft = revision["revised_draft"]

            round_data = {
                "round_num": round_num,
                "review": {k: v for k, v in consolidated.items() if not k.startswith("_")},
                "revision": {
                    "accepted_count": revision.get("accepted_count", 0),
                    "rebutted_count": revision.get("rebutted_count", 0),
                    "deferred_count": revision.get("deferred_count", 0),
                    "response_to_reviewers": revision.get("response_to_reviewers", []),
                    "triage": revision.get("triage", []),
                },
            }
            all_rounds_data.append(round_data)

            _emit_sse(upload_id, "revision_complete", {
                "round": round_num,
                "accepted": revision.get("accepted_count", 0),
                "rebutted": revision.get("rebutted_count", 0),
            }, session_id)

            # Save progress after each round
            conn = get_connection()
            conn.execute(
                """UPDATE paper_uploads SET review_rounds = ?, current_draft = ?,
                   score_progression = ?, updated_at = ? WHERE upload_id = ?""",
                (json.dumps(all_rounds_data, default=str), current_draft,
                 json.dumps(score_progression), datetime.now().isoformat(), upload_id),
            )
            conn.commit()

            # Convergence check
            if len(score_progression) >= 2:
                delta = score_progression[-1] - score_progression[-2]
                if delta < 0.5 and score_progression[-1] >= 6.0:
                    _emit_sse(upload_id, "converged", {
                        "round": round_num, "score": score_progression[-1],
                    }, session_id)
                    break

        # Finalize
        conn = get_connection()
        conn.execute(
            "UPDATE paper_uploads SET status = ?, updated_at = ? WHERE upload_id = ?",
            ("review_complete", datetime.now().isoformat(), upload_id),
        )
        conn.commit()

        tm.update_task(task_id, status=TaskStatus.COMPLETED, progress=100,
                       message="Review complete",
                       result={
                           "rounds_completed": len(all_rounds_data),
                           "initial_score": score_progression[0] if score_progression else 0,
                           "final_score": score_progression[-1] if score_progression else 0,
                           "final_decision": all_consolidations[-1].get("final_decision", "unknown") if all_consolidations else "unknown",
                       })
        _emit_sse(upload_id, "complete", {
            "initial_score": score_progression[0] if score_progression else 0,
            "final_score": score_progression[-1] if score_progression else 0,
        }, session_id)

    except Exception as e:
        logger.exception("Review pipeline failed")
        conn = get_connection()
        conn.execute(
            "UPDATE paper_uploads SET status = ?, error = ?, updated_at = ? WHERE upload_id = ?",
            ("review_failed", str(e), datetime.now().isoformat(), upload_id),
        )
        conn.commit()
        tm.update_task(task_id, status=TaskStatus.FAILED, error=str(e))
        _emit_sse(upload_id, "error", {"message": str(e)}, session_id)


# ── Status & Data Endpoints ────────────────────────────────────────


@paper_rehab_bp.route("/paper-lab/<upload_id>/status", methods=["GET"])
def get_status(upload_id):
    """Get current review status + scores."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    metadata = upload["metadata"] or {}

    return jsonify({
        "success": True,
        "data": {
            **_summarize_upload(upload),
            "section_count": len(upload["sections"]),
            "sections": [s.get("name", "") for s in upload["sections"]],
            "word_count": metadata.get("word_count", 0),
            "reference_count": len(upload["raw_references"]),
            "review_config": upload["review_config"],
            "error": upload["error"],
            "parse_warnings": metadata.get("parse_warnings", []),
            "parse_quality_breakdown": metadata.get("parse_quality", {}),
            "tables": metadata.get("tables", []),
            "figures": metadata.get("figures", []),
            "formulas": metadata.get("formulas", []),
        },
    })


@paper_rehab_bp.route("/paper-lab/<upload_id>/rounds", methods=["GET"])
def get_rounds(upload_id):
    """Fetch all review/revision round data."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    return jsonify({
        "success": True,
        "data": {
            "upload_id": upload_id,
            "rounds": upload["review_rounds"],
            "score_progression": upload["score_progression"],
            "source_audit": upload["source_audit"],
            "reviewers": upload["reviewers"],
            "authors": upload["authors"],
        },
    })


@paper_rehab_bp.route("/paper-lab/<upload_id>/draft", methods=["GET"])
def get_draft(upload_id):
    """Get the current (latest revised) draft text."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    return jsonify({
        "success": True,
        "data": {
            "upload_id": upload_id,
            "title": upload["title"],
            "draft": upload["current_draft"],
            "word_count": len(upload["current_draft"].split()),
            "status": upload["status"],
        },
    })


@paper_rehab_bp.route("/paper-lab/<upload_id>/stream", methods=["GET"])
def stream_progress(upload_id):
    """SSE stream for live review/revision progress."""
    upload = _load_upload(upload_id) or {}
    requested_session_id = (request.args.get("session_id") or _current_review_session_id(upload)).strip()

    def generate():
        last_idx = 0
        timeout = 600  # 10 minutes max
        start = time.time()
        session_id = requested_session_id

        yield f"data: {json.dumps({'type': 'connected', 'upload_id': upload_id, 'session_id': session_id})}\n\n"

        while time.time() - start < timeout:
            events = _sse_events.get(upload_id, [])
            filtered_events = [evt for evt in events if not session_id or evt.get("session_id") == session_id]
            if last_idx < len(filtered_events):
                for evt in filtered_events[last_idx:]:
                    yield f"event: {evt['type']}\ndata: {json.dumps(evt['data'], default=str)}\n\n"
                last_idx = len(filtered_events)

                # Check if complete
                if filtered_events and filtered_events[-1]["type"] in ("complete", "error"):
                    break
            time.sleep(1)

        yield f"data: {json.dumps({'type': 'stream_end'})}\n\n"

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


@paper_rehab_bp.route("/paper-lab/uploads", methods=["GET"])
def list_uploads():
    """List all paper uploads."""
    conn = get_connection()
    rows = conn.execute(
        "SELECT upload_id FROM paper_uploads ORDER BY created_at DESC"
    ).fetchall()

    uploads = []
    for row in rows:
        upload = _load_upload(row["upload_id"])
        if upload:
            uploads.append(_summarize_upload(upload))

    return jsonify({"success": True, "data": uploads})


# ── Specialist Review ───────────────────────────────────────────────


@paper_rehab_bp.route("/paper-lab/<upload_id>/specialist-review", methods=["POST"])
def run_specialist_review(upload_id):
    """
    Run specialist domain review against paper upload content.
    Body: { domains?: string[], strictness?: float, target?: "draft", model?: string }
    """
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    data = request.get_json(silent=True) or {}
    domains = data.get("domains")
    target = data.get("target", "draft")
    model = data.get("model", "")

    try:
        strictness = float(data.get("strictness", 0.7))
    except (TypeError, ValueError):
        strictness = 0.7
    strictness = min(max(strictness, 0.0), 1.0)

    if target != "draft":
        return jsonify({
            "success": False,
            "error": f"Unsupported target '{target}' for Paper Lab specialist review. Use 'draft'.",
        }), 400

    content = (upload.get("current_draft") or upload.get("full_text") or "").strip()
    if not content:
        return jsonify({"success": False, "error": "No draft content available for specialist review"}), 400

    tm = TaskManager()
    task_id = tm.create_task(task_type="paper_specialist_review", metadata={"upload_id": upload_id, "target": target})

    def _run_specialist_review_task():
        try:
            tm.update_task(
                task_id,
                status=TaskStatus.PROCESSING,
                progress=10,
                message="Running specialist domain reviews...",
            )

            from ..services.ais.specialist_review import SpecialistReviewService
            reviewer = SpecialistReviewService()
            results = reviewer.review(content, domains=domains, strictness=strictness, model=model)
            review_data = [r.to_dict() for r in results]

            payload = {
                "upload_id": upload_id,
                "target": target,
                "domains": domains or [r.get("domain") for r in review_data],
                "strictness": strictness,
                "reviews": review_data,
                "domain_count": len(results),
                "total_findings": sum(len(r.findings) for r in results),
                "reviewed_at": datetime.now().isoformat(),
            }

            latest = _load_upload(upload_id)
            if latest:
                meta = latest.get("metadata", {}) or {}
                meta["specialist_review"] = payload
                conn = get_connection()
                # Bug #5 fix: also update status so frontend can track the
                # specialist-review lifecycle across refreshes. Only advance
                # status forward (don't regress from review_complete).
                current_status = latest.get("status", "")
                new_status = current_status
                if current_status in ("review_complete", "gap_filled"):
                    new_status = "specialist_complete"
                conn.execute(
                    "UPDATE paper_uploads SET metadata = ?, status = ?, updated_at = ? WHERE upload_id = ?",
                    (json.dumps(meta), new_status, datetime.now().isoformat(), upload_id),
                )
                conn.commit()

            tm.update_task(task_id, status=TaskStatus.PROCESSING, progress=100, message="Specialist review complete")
            tm.update_task(task_id, status=TaskStatus.COMPLETED, progress=100, result=payload)
        except Exception as e:
            logger.exception("Paper specialist review failed")
            tm.update_task(task_id, status=TaskStatus.FAILED, error=str(e))

    thread = threading.Thread(target=_run_specialist_review_task, daemon=True)
    thread.start()

    return jsonify({
        "success": True,
        "data": {
            "upload_id": upload_id,
            "task_id": task_id,
            "message": "Specialist review started.",
        },
    }), 202


@paper_rehab_bp.route("/paper-lab/<upload_id>/specialist-review", methods=["GET"])
def get_specialist_review(upload_id):
    """Return latest specialist review snapshot for a paper upload."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    specialist = (upload.get("metadata") or {}).get("specialist_review")
    if not specialist:
        return jsonify({"success": False, "error": "No specialist review found"}), 404

    return jsonify({"success": True, "data": specialist})


# ── Gap Fill & Novelty Boost ───────────────────────────────────────


@paper_rehab_bp.route("/paper-lab/<upload_id>/fill-gaps", methods=["POST"])
def fill_gaps(upload_id):
    """
    Search for missing references and evidence based on reviewer feedback.
    Runs gap-fill + novelty boost, updates the draft.
    Body: { "live": true }
    """
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    if upload["status"] != "review_complete":
        return jsonify({"success": False, "error": f"Review not complete: status is '{upload['status']}'"}), 400

    data = request.get_json() or {}
    use_live = data.get("live", False)

    try:
        import sys
        sys.path.insert(0, str(Path(__file__).parent.parent.parent))
        from cli_test_paper_rehab import test_gap_fill, test_novelty_boost
        from app.services.ais.paper_parser import ParsedPaper
        from app.models.ais_models import PaperSection

        # Reconstruct ParsedPaper
        sections = [PaperSection.from_dict(s) for s in upload["sections"]]
        parsed_paper = ParsedPaper(
            source_path=upload.get("source_filename", ""),
            title=upload["title"],
            language=upload["language"],
            detected_field=upload["detected_field"],
            sections=sections,
            raw_references=upload["raw_references"],
            metadata=upload["metadata"],
            full_text=upload["full_text"],
        )

        all_consolidations = [rd.get("review", {}) for rd in upload["review_rounds"]]
        current_draft = upload["current_draft"]

        # Gap fill
        gap_result = test_gap_fill(parsed_paper, all_consolidations, current_draft, use_live)
        found_papers = gap_result.get("_found_papers", [])
        gap_fill_text = gap_result.get("_gap_fill_text", "")

        # Novelty boost
        novelty_result = test_novelty_boost(
            parsed_paper, all_consolidations, current_draft,
            found_papers, gap_fill_text, use_live,
        )
        novelty_boost = novelty_result.get("_novelty_boost", "")

        # Update draft
        if gap_fill_text:
            current_draft += f"\n\n## Additional Literature (Gap Fill)\n\n{gap_fill_text}"
        if novelty_boost:
            current_draft += f"\n\n## Novelty Reframing (Suggested Additions)\n\n{novelty_boost}"

        # Save to DB
        conn = get_connection()
        conn.execute(
            "UPDATE paper_uploads SET current_draft = ?, status = ?, updated_at = ? WHERE upload_id = ?",
            (current_draft, "gap_filled", datetime.now().isoformat(), upload_id),
        )
        conn.commit()

        return jsonify({
            "success": True,
            "data": {
                "upload_id": upload_id,
                "gap_fill": {k: v for k, v in gap_result.items() if not k.startswith("_")},
                "novelty": {k: v for k, v in novelty_result.items() if not k.startswith("_")},
                "papers_found": len(found_papers),
                "found_papers": [
                    {"title": p.get("title", ""), "doi": p.get("doi", ""), "year": p.get("year", "")}
                    for p in found_papers[:10]
                ],
                "draft_updated": True,
            },
        })
    except Exception as e:
        logger.exception("Gap fill failed")
        return jsonify({"success": False, "error": str(e)}), 500


@paper_rehab_bp.route("/paper-lab/<upload_id>/export-docx", methods=["GET"])
def export_docx(upload_id):
    """Export the current draft as a .docx file download."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = DocxDocument()
        title_para = doc.add_heading(upload["title"], level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        draft = upload["current_draft"]
        lines = draft.split("\n")
        current_text = []

        for line in lines:
            stripped = line.strip()
            if stripped.startswith("## "):
                if current_text:
                    doc.add_paragraph("\n".join(current_text))
                    current_text = []
                doc.add_heading(stripped.lstrip("# ").strip(), level=2)
            elif stripped.startswith("# "):
                if current_text:
                    doc.add_paragraph("\n".join(current_text))
                    current_text = []
                doc.add_heading(stripped.lstrip("# ").strip(), level=1)
            elif stripped:
                current_text.append(stripped)

        if current_text:
            doc.add_paragraph("\n".join(current_text))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        # ASCII-safe filename (Vietnamese/CJK chars can't be in Content-Disposition)
        safe_title = upload["title"].encode("ascii", "ignore").decode("ascii")
        safe_title = re.sub(r'[^\w\s-]', '', safe_title)[:50].strip() or "paper_draft"
        filename = f"{safe_title}_{upload_id}.docx"

        return Response(
            buf.getvalue(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except ImportError:
        return jsonify({"success": False, "error": "python-docx not installed"}), 500
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@paper_rehab_bp.route("/paper-lab/<upload_id>/response-to-reviewers", methods=["GET"])
def get_response_to_reviewers(upload_id):
    """Generate and return Response to Reviewers as markdown or .docx download."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    fmt = request.args.get("format", "json")  # json, markdown, docx

    try:
        from ..services.ais.response_generator import (
            generate_response_to_reviewers,
            response_to_docx,
        )

        result = generate_response_to_reviewers(
            paper_title=upload["title"],
            review_rounds=upload["review_rounds"],
            score_progression=upload["score_progression"],
            source_audit=upload["source_audit"],
        )

        if fmt == "docx":
            docx_bytes = response_to_docx(result["markdown"], upload["title"])
            return Response(
                docx_bytes,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="response_to_reviewers_{upload_id}.docx"'},
            )
        elif fmt == "markdown":
            return Response(result["markdown"], mimetype="text/markdown",
                            headers={"Content-Disposition": f'attachment; filename="response_to_reviewers_{upload_id}.md"'})
        else:
            return jsonify({"success": True, "data": result})

    except Exception as e:
        logger.exception("Response generation failed")
        return jsonify({"success": False, "error": str(e)}), 500


@paper_rehab_bp.route("/paper-lab/<upload_id>/rewrite-instructions", methods=["GET"])
def get_rewrite_instructions(upload_id):
    """Generate a comprehensive rewrite instruction document for Claude Opus/Sonnet.
    Includes: original draft, all reviewer feedback, triage decisions, and
    detailed instructions for rewriting with diagrams/graphs."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    title = upload.get("title", "Untitled")
    field = upload.get("detected_field", upload.get("field", "general"))
    language = upload.get("language", "en")
    current_draft = upload.get("current_draft", "")
    review_rounds = upload.get("review_rounds", [])
    score_progression = upload.get("score_progression", [])
    source_audit = upload.get("source_audit", {})

    # Build reviewer feedback summary
    reviewer_feedback = []
    for i, rnd in enumerate(review_rounds):
        round_reviews = rnd.get("reviews", [])
        consolidated = rnd.get("consolidated", {})
        revision = rnd.get("revision", {})
        triage = revision.get("triage", [])

        reviewer_feedback.append(f"\n### Round {i + 1}")
        if consolidated:
            reviewer_feedback.append(f"**Average Score:** {consolidated.get('avg_overall_score', '?')}/10")
            reviewer_feedback.append(f"**Decision:** {consolidated.get('final_decision', '?')}")

        # Individual reviews
        for rev in round_reviews:
            reviewer_feedback.append(f"\n**{rev.get('reviewer', 'Reviewer')}** (Score: {rev.get('score', '?')}/10)")
            for w in rev.get("weaknesses", []):
                if isinstance(w, dict):
                    reviewer_feedback.append(f"- [{w.get('severity', 'major')}] {w.get('section', '?')}: {w.get('text', w.get('description', ''))}")
                else:
                    reviewer_feedback.append(f"- {w}")

        # Consolidated weaknesses
        all_weak = consolidated.get("all_weaknesses", consolidated.get("top_weaknesses", []))
        if all_weak and not round_reviews:
            reviewer_feedback.append("\n**Consolidated Weaknesses:**")
            for w in all_weak:
                if isinstance(w, dict):
                    reviewer_feedback.append(f"- [{w.get('severity', 'major')}] {w.get('section', '?')}: {w.get('text', w.get('description', ''))}")
                else:
                    reviewer_feedback.append(f"- {w}")

        # Triage decisions
        if triage:
            accepted = [t for t in triage if isinstance(t, dict) and t.get("action") == "accept"]
            rebutted = [t for t in triage if isinstance(t, dict) and t.get("action") == "rebut"]
            deferred = [t for t in triage if isinstance(t, dict) and t.get("action") == "defer"]
            reviewer_feedback.append(f"\n**Author Triage:** {len(accepted)} accepted, {len(rebutted)} rebutted, {len(deferred)} deferred")

    feedback_text = "\n".join(reviewer_feedback) if reviewer_feedback else "(No review data available — please run Start Review first)"

    # Source audit summary
    audit_text = ""
    if source_audit:
        verified = source_audit.get("verified", [])
        unverified = source_audit.get("unverified", [])
        audit_text = f"\n**Verified references:** {len(verified)}\n**Unverified references:** {len(unverified)}"
        if unverified:
            audit_text += "\n\nUnverified references to check or replace:\n"
            for ref in unverified[:10]:
                if isinstance(ref, dict):
                    audit_text += f"- {ref.get('title', ref.get('raw', str(ref)))}\n"
                else:
                    audit_text += f"- {ref}\n"

    # Score progression
    score_text = ""
    if score_progression:
        score_text = "Score progression: " + " → ".join(f"R{i+1}: {s:.1f}" for i, s in enumerate(score_progression))

    instructions = f"""# Rewrite Instructions for: {title}

> **Generated by Parallax Paper Lab** — {datetime.now().strftime('%Y-%m-%d %H:%M')}
> Field: {field} | Language: {language}
> {score_text}

---

## Your Task

You are an expert academic writer and researcher. Rewrite the draft paper below based on peer reviewer feedback. Your rewrite should:

1. **Address all accepted reviewer points** (listed below)
2. **Improve academic English** — formal tone, precise terminology, clear argumentation
3. **Strengthen the methodology section** with more detail on experimental setup, parameters, and validation
4. **Add diagrams/figures descriptions** — for each key concept, describe a figure that should be created:
   - Include figure captions in the format: `[Figure N: Description of what the figure shows]`
   - Suggest: experimental setup diagrams, data flow charts, comparison graphs, SEM/TEM illustrations
5. **Add proper citations** — use IEEE/APA format, cite real papers from the field
6. **Improve the abstract** — concise, include method, key results, and contribution
7. **Add a clear conclusion** with future work directions

## Quality Criteria
- Target score: 7.0+ / 10.0 (current: {f'{score_progression[-1]:.1f}' if score_progression else 'N/A'})
- Word count: aim for 4000-6000 words (current: {len(current_draft.split()) if current_draft else 0})
- Must include: Abstract, Introduction, Background, Methodology, Results, Discussion, Conclusion, References

---

## Peer Reviewer Feedback

{feedback_text}

{audit_text}

---

## Original Draft (to rewrite)

{current_draft}

---

## Output Format

Please produce:
1. The **complete rewritten paper** in academic format (markdown with proper headings)
2. A **list of figures/diagrams** that should be created, with detailed descriptions
3. A **changelog** summarizing what you changed and why

Start your rewrite now.
"""

    fmt = request.args.get("format", "markdown")
    if fmt == "json":
        return jsonify({
            "success": True,
            "data": {
                "instructions": instructions,
                "title": title,
                "word_count": len(instructions.split()),
                "has_reviews": len(review_rounds) > 0,
                "score_progression": score_progression,
            },
        })
    else:
        safe_title = re.sub(r'[^\w\s-]', '', title)[:50].strip().replace(' ', '_') or 'draft'
        return Response(
            instructions,
            mimetype="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="rewrite_instructions_{safe_title}.md"'},
        )


# ── Visualization Handler ──────────────────────────────────────────────────


def _get_viz_cache(upload_id: str) -> dict:
    """Load cached visualization results from upload metadata."""
    upload = _load_upload(upload_id)
    if not upload:
        return {}
    return (upload.get("metadata") or {}).get("visualizations", {})


def _save_viz_cache(upload_id: str, key: str, data: dict):
    """Persist a visualization result into upload metadata."""
    upload = _load_upload(upload_id)
    if not upload:
        return
    meta = upload.get("metadata") or {}
    viz = meta.get("visualizations", {})
    viz[key] = data
    meta["visualizations"] = viz
    conn = get_connection()
    conn.execute(
        "UPDATE paper_uploads SET metadata = ?, updated_at = ? WHERE upload_id = ?",
        (json.dumps(meta), datetime.now().isoformat(), upload_id),
    )
    conn.commit()


def _artifact_provenance(upload: dict, source: str) -> dict:
    return {
        "generated_by": source,
        "generated_at": datetime.now().isoformat(),
        "derived_from_upload_version": upload.get("updated_at") or upload.get("created_at") or "",
        "upload_status": upload.get("status", ""),
    }


@paper_rehab_bp.route("/paper-lab/<upload_id>/analyze-figures", methods=["POST"])
def analyze_figures(upload_id):
    """
    Detect figure references, generate Python reconstruction code, flag data gaps.
    Body: {} (no required fields)
    Returns: task_id — poll /visualizations for results.
    """
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    full_text = upload.get("current_draft") or upload.get("full_text") or ""
    sections = upload.get("sections") or []

    tm = TaskManager()
    task_id = tm.create_task("viz_figures", metadata={"upload_id": upload_id})

    def _task():
        try:
            from ..services.ais.visualization_service import analyze_figures as _analyze
            tm.update_task(task_id, status=TaskStatus.PROCESSING, progress=10, message="Analysing figures...")
            result = _analyze(full_text, sections)
            _save_viz_cache(upload_id, "figures", result)
            tm.update_task(task_id, status=TaskStatus.COMPLETED, progress=100, result=result)
        except Exception as e:
            logger.exception("analyze_figures task failed")
            tm.update_task(task_id, status=TaskStatus.FAILED, error=str(e))

    threading.Thread(target=_task, daemon=True).start()
    return jsonify({"success": True, "data": {"task_id": task_id}}), 202


@paper_rehab_bp.route("/paper-lab/<upload_id>/analyze-tables", methods=["POST"])
def analyze_tables(upload_id):
    """
    Extract tables, flag statistical errors, propose corrected data.
    Returns: task_id — poll /visualizations for results.
    """
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    full_text = upload.get("current_draft") or upload.get("full_text") or ""
    sections = upload.get("sections") or []

    tm = TaskManager()
    task_id = tm.create_task("viz_tables", metadata={"upload_id": upload_id})

    def _task():
        try:
            from ..services.ais.visualization_service import analyze_tables as _analyze
            tm.update_task(task_id, status=TaskStatus.PROCESSING, progress=10, message="Analysing tables...")
            result = _analyze(full_text, sections)
            _save_viz_cache(upload_id, "tables", result)
            tm.update_task(task_id, status=TaskStatus.COMPLETED, progress=100, result=result)
        except Exception as e:
            logger.exception("analyze_tables task failed")
            tm.update_task(task_id, status=TaskStatus.FAILED, error=str(e))

    threading.Thread(target=_task, daemon=True).start()
    return jsonify({"success": True, "data": {"task_id": task_id}}), 202


@paper_rehab_bp.route("/paper-lab/<upload_id>/generate-diagram", methods=["POST"])
def generate_diagram(upload_id):
    """
    Generate a Mermaid.js diagram using Gemini 2.0 Flash.
    Body: { "diagram_type": "flowchart|mindmap|sequence|timeline|quadrant|infographic" }
    Synchronous — returns the diagram directly.
    """
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    data = request.get_json(silent=True) or {}
    diagram_type = data.get("diagram_type", "flowchart")

    # Extract abstract from sections if available
    sections = upload.get("sections") or []
    abstract = ""
    for s in sections:
        if "abstract" in (s.get("name") or "").lower():
            abstract = s.get("text") or s.get("content") or ""
            break
    if not abstract:
        abstract = (upload.get("full_text") or "")[:1200]

    # Extract key findings from last review round if available
    key_findings = []
    rounds = upload.get("review_rounds") or []
    if rounds:
        last = rounds[-1]
        review = last.get("review") or {}
        for weakness in (review.get("all_weaknesses") or [])[:5]:
            if isinstance(weakness, dict):
                key_findings.append(weakness.get("text", ""))
            elif isinstance(weakness, str):
                key_findings.append(weakness)

    try:
        from ..services.ais.visualization_service import generate_diagram as _gen
        result = _gen(
            title=upload.get("title", "Untitled"),
            abstract=abstract,
            key_findings=key_findings,
            diagram_type=diagram_type,
        )
        # Cache the result
        _save_viz_cache(upload_id, f"diagram_{diagram_type}", result)
        return jsonify({"success": True, "data": result})
    except Exception as e:
        logger.exception("generate_diagram failed")
        return jsonify({"success": False, "error": str(e)}), 500


@paper_rehab_bp.route("/paper-lab/<upload_id>/deep-analysis", methods=["POST"])
def deep_analysis(upload_id):
    """
    Deep cross-data analysis using Gemini 2.0 Flash Thinking.
    Proposes simulations, cross-dataset strategies, improved statements.
    Returns: task_id — poll /visualizations for results.
    """
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    full_text = upload.get("current_draft") or upload.get("full_text") or ""
    sections = upload.get("sections") or []
    review_rounds = upload.get("review_rounds") or []

    tm = TaskManager()
    task_id = tm.create_task("viz_deep_analysis", metadata={"upload_id": upload_id})

    def _task():
        try:
            from ..services.ais.visualization_service import deep_analysis as _analyze
            tm.update_task(task_id, status=TaskStatus.PROCESSING, progress=5,
                           message="Running deep thinking analysis (this may take 30-60s)...")
            result = _analyze(full_text, sections, review_rounds)
            _save_viz_cache(upload_id, "deep_analysis", result)
            tm.update_task(task_id, status=TaskStatus.COMPLETED, progress=100, result=result)
        except Exception as e:
            logger.exception("deep_analysis task failed")
            tm.update_task(task_id, status=TaskStatus.FAILED, error=str(e))

    threading.Thread(target=_task, daemon=True).start()
    return jsonify({"success": True, "data": {"task_id": task_id}}), 202


@paper_rehab_bp.route("/paper-lab/<upload_id>/visualizations", methods=["GET"])
def get_visualizations(upload_id):
    """
    Return all cached visualization results for a paper upload.
    Keys: figures, tables, diagram_*, deep_analysis, rendered_figures, figure_audit
    """
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    viz = _get_viz_cache(upload_id)
    return jsonify({"success": True, "data": viz})


@paper_rehab_bp.route("/paper-lab/<upload_id>/visualization-plan", methods=["POST"])
def visualization_plan(upload_id):
    """Generate structured visualization recommendations from manuscript + reviews."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    from ..services.ais.paper_orchestra_service import PaperOrchestraService

    service = PaperOrchestraService()
    plan = service.build_visualization_plan(upload)
    return jsonify({"success": True, "data": plan})


@paper_rehab_bp.route("/paper-lab/<upload_id>/visualization-artifacts", methods=["GET"])
def list_visualization_artifacts(upload_id):
    """List persisted visualization artifacts for a manuscript."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    from ..services.ais.visualization_artifact_service import VisualizationArtifactService

    service = VisualizationArtifactService()
    return jsonify({"success": True, "data": service.list_for_upload(upload_id)})


@paper_rehab_bp.route("/paper-lab/<upload_id>/artifacts", methods=["POST"])
def create_visualization_artifact(upload_id):
    """Create a persisted visualization artifact."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    payload = request.get_json(silent=True) or {}
    artifact_type = payload.get("type", "chart")
    intent = payload.get("intent", "summarize")
    title = payload.get("title") or "Untitled Artifact"
    from ..services.ais.visualization_artifact_service import VisualizationArtifactService

    service = VisualizationArtifactService()
    artifact = service.create(
        upload_id=upload_id,
        artifact_type=artifact_type,
        intent=intent,
        title=title,
        payload=payload.get("payload", {}),
        audit=payload.get("audit", {}),
        provenance={**_artifact_provenance(upload, "manual_create"), **(payload.get("provenance") or {})},
        status=payload.get("status", "draft"),
    )
    return jsonify({"success": True, "data": artifact}), 201


@paper_rehab_bp.route("/paper-lab/<upload_id>/artifacts/<artifact_id>", methods=["GET"])
def get_visualization_artifact(upload_id, artifact_id):
    """Get one persisted visualization artifact."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    from ..services.ais.visualization_artifact_service import VisualizationArtifactService

    artifact = VisualizationArtifactService().get(artifact_id)
    if not artifact or artifact.get("upload_id") != upload_id:
        return jsonify({"success": False, "error": "Artifact not found"}), 404
    return jsonify({"success": True, "data": artifact})


@paper_rehab_bp.route("/paper-lab/<upload_id>/artifacts/<artifact_id>", methods=["PUT"])
def update_visualization_artifact(upload_id, artifact_id):
    """Update a persisted visualization artifact."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    from ..services.ais.visualization_artifact_service import VisualizationArtifactService

    body = request.get_json(silent=True) or {}
    service = VisualizationArtifactService()
    artifact = service.get(artifact_id)
    if not artifact or artifact.get("upload_id") != upload_id:
        return jsonify({"success": False, "error": "Artifact not found"}), 404

    updated = service.update(
        artifact_id,
        title=body.get("title"),
        status=body.get("status"),
        payload_patch=body.get("payload"),
        audit_patch=body.get("audit"),
        provenance_patch=body.get("provenance"),
        increment_version=bool(body.get("increment_version", True)),
    )
    return jsonify({"success": True, "data": updated})


@paper_rehab_bp.route("/paper-lab/<upload_id>/artifacts/<artifact_id>/render", methods=["POST"])
def render_visualization_artifact(upload_id, artifact_id):
    """Render an artifact into a browser-oriented representation."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    from ..services.ais.visualization_artifact_service import VisualizationArtifactService
    service = VisualizationArtifactService()
    artifact = service.get(artifact_id)
    if not artifact or artifact.get("upload_id") != upload_id:
        return jsonify({"success": False, "error": "Artifact not found"}), 404

    payload = artifact.get("payload", {})
    rendering = payload.get("rendering", {})
    artifact_type = artifact.get("type", "chart")
    default_formats = {
        "graphical_abstract": ["html", "json"],
        "slide": ["json"],
        "poster_panel": ["json"],
        "chart": ["json", "svg", "png"],
        "diagram": ["json", "svg"],
        "table": ["json", "csv"],
    }.get(artifact_type, ["json"])
    export_formats = list(dict.fromkeys((payload.get("export_formats") or []) + default_formats))
    inferred_engine = payload.get("recommended_engine") or rendering.get("engine") or (
        "html" if artifact_type in {"graphical_abstract", "slide", "poster_panel"} else "vega-lite"
    )
    inferred_spec = rendering.get("spec")
    if inferred_spec is None:
        if artifact_type == "slide":
            inferred_spec = {"slides": payload.get("slides", [])}
        elif artifact_type == "poster_panel":
            inferred_spec = {"panels": payload.get("panels", [])}
        else:
            inferred_spec = payload.get("spec") or payload.get("content_description") or payload
    updated = service.update(
        artifact_id,
        status="ready" if not payload.get("assumptions") else "needs_input",
        payload_patch={
            "rendering": {
                **rendering,
                "engine": inferred_engine,
                "spec": inferred_spec,
                "rendered_at": datetime.now().isoformat(),
            },
            "export_formats": export_formats,
        },
        provenance_patch=_artifact_provenance(upload, "artifact_render"),
        increment_version=True,
    )
    return jsonify({"success": True, "data": updated})


@paper_rehab_bp.route("/paper-lab/<upload_id>/artifacts/<artifact_id>/audit", methods=["POST"])
def audit_visualization_artifact(upload_id, artifact_id):
    """Audit an artifact and persist readiness signals."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    from ..services.ais.visualization_artifact_service import VisualizationArtifactService
    service = VisualizationArtifactService()
    artifact = service.get(artifact_id)
    if not artifact or artifact.get("upload_id") != upload_id:
        return jsonify({"success": False, "error": "Artifact not found"}), 404

    payload = artifact.get("payload", {})
    assumptions = payload.get("assumptions") or []
    rendering = payload.get("rendering") or {}
    has_rendering = rendering.get("spec") is not None
    issues = []
    if assumptions:
        issues.append("Resolve inferred data assumptions before export.")
    if artifact.get("type") in {"chart", "diagram", "graphical_abstract"} and not has_rendering:
        issues.append("Render the artifact before export.")
    consistency_status = "pass"
    if issues:
        consistency_status = "fail" if not has_rendering else "warn"
    audit = {
        "confidence": 0.93 if not issues else (0.72 if has_rendering else 0.44),
        "issues": issues,
        "consistency_status": consistency_status,
        "ready": not issues,
    }
    updated = service.update(
        artifact_id,
        status="ready" if audit["ready"] else "needs_input",
        audit_patch=audit,
        provenance_patch=_artifact_provenance(upload, "artifact_audit"),
        increment_version=True,
    )
    return jsonify({"success": True, "data": updated})


@paper_rehab_bp.route("/paper-lab/<upload_id>/artifacts/<artifact_id>/export", methods=["POST"])
def export_visualization_artifact(upload_id, artifact_id):
    """Return an export-ready artifact package."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    from ..services.ais.visualization_artifact_service import VisualizationArtifactService
    artifact = VisualizationArtifactService().get(artifact_id)
    if not artifact or artifact.get("upload_id") != upload_id:
        return jsonify({"success": False, "error": "Artifact not found"}), 404

    audit = artifact.get("audit") or {}
    payload = artifact.get("payload", {})
    blocked_by = []
    if payload.get("assumptions"):
        blocked_by.extend(payload.get("assumptions", []))
    blocked_by.extend(audit.get("issues") or [])
    ready = not blocked_by and bool(audit.get("ready", False))
    payload = artifact.get("payload", {})
    package = _artifact_export_bundle(artifact)
    return jsonify({
        "success": True,
        "data": {
            "artifact": artifact,
            "ready": ready,
            "blocked_by": blocked_by if not ready else [],
            "package": package,
        },
    })


@paper_rehab_bp.route("/paper-lab/<upload_id>/render-figures", methods=["POST"])
def render_figures(upload_id):
    """
    Generate browser-renderable Vega-Lite specs from figure analysis results.
    Requires analyze-figures to have been run first.
    Returns rendered specs immediately (no async task needed).
    """
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    viz = _get_viz_cache(upload_id)
    figure_analysis = viz.get("figures")
    if not figure_analysis:
        return jsonify({"success": False, "error": "Run figure analysis first (POST /analyze-figures)"}), 400

    try:
        from ..services.ais.scientific_viz import render_figures as _render
        from ..services.ais.visualization_artifact_service import VisualizationArtifactService
        rendered = _render(figure_analysis)
        _save_viz_cache(upload_id, "rendered_figures", {"figures": rendered, "count": len(rendered)})
        artifact_service = VisualizationArtifactService()
        for fig in rendered:
            artifact_service.create_or_replace_by_title(
                upload_id=upload_id,
                artifact_type="chart",
                intent="reconstruct",
                title=fig.get("ref") or fig.get("title") or "Rendered Figure",
                payload={
                    "source_refs": [fig.get("ref", "")],
                    "rendering": {
                        "engine": "vega-lite",
                        "spec": fig.get("vega_lite_spec"),
                    },
                    "export_formats": ["json", "svg", "png"],
                    "assumptions": list(fig.get("data_requirements", []) or []),
                    "issues": list(fig.get("issues", []) or []),
                },
                provenance=_artifact_provenance(upload, "render_figures"),
                status="needs_input" if fig.get("data_requirements") else "ready",
            )
        return jsonify({"success": True, "data": {"figures": rendered, "count": len(rendered)}})
    except Exception as e:
        logger.exception("render_figures failed")
        return jsonify({"success": False, "error": str(e)}), 500


@paper_rehab_bp.route("/paper-lab/<upload_id>/audit-figures", methods=["POST"])
def audit_figures_endpoint(upload_id):
    """
    Audit figures against Rougier's Ten Simple Rules for Better Figures.
    Requires analyze-figures to have been run first.
    Returns audit results immediately.
    """
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    viz = _get_viz_cache(upload_id)
    figure_analysis = viz.get("figures")
    if not figure_analysis:
        return jsonify({"success": False, "error": "Run figure analysis first (POST /analyze-figures)"}), 400

    full_text = upload.get("full_text", "")

    try:
        from ..services.ais.scientific_viz import audit_figures as _audit
        from ..services.ais.visualization_artifact_service import VisualizationArtifactService
        audit_result = _audit(figure_analysis, full_text)
        _save_viz_cache(upload_id, "figure_audit", audit_result)
        artifact_service = VisualizationArtifactService()
        for entry in audit_result.get("figures", []):
            artifact_service.create_or_replace_by_title(
                upload_id=upload_id,
                artifact_type="chart",
                intent="reconstruct",
                title=entry.get("ref") or "Audited Figure",
                audit={
                    "confidence": float(entry.get("score", 0)) / 10.0,
                    "issues": [check.get("note", "") for check in entry.get("checks", []) if check.get("status") != "pass"],
                    "consistency_status": "pass" if float(entry.get("score", 0)) >= 8 else "warn",
                    "ready": float(entry.get("score", 0)) >= 8,
                },
                provenance=_artifact_provenance(upload, "audit_figures"),
                status="ready" if float(entry.get("score", 0)) >= 8 else "needs_input",
            )
        return jsonify({"success": True, "data": audit_result})
    except Exception as e:
        logger.exception("audit_figures failed")
        return jsonify({"success": False, "error": str(e)}), 500


@paper_rehab_bp.route("/paper-lab/compare", methods=["POST"])
def compare_papers_endpoint():
    """
    POST /paper-lab/compare
    Body: {"upload_ids": ["uuid-1", "uuid-2"]}
    Spawns background task `compare_manuscripts` and returns a task id.
    """
    data = request.json or {}
    upload_ids = data.get("upload_ids")
    
    if not upload_ids or not isinstance(upload_ids, list) or len(upload_ids) < 2:
        return jsonify({"success": False, "error": "Provide at least two upload_ids for comparative analysis"}), 400
        
    documents = []
    for uid in upload_ids:
        up = _load_upload(uid)
        if up:
            documents.append({
                "title": up.get("detected_title", f"Document {uid}"),
                "text": up.get("full_text", "")
            })
            
    if len(documents) < 2:
         return jsonify({"success": False, "error": "Could not locate valid text for the requested documents"}), 404

    from ..services.ais.visualization_service import compare_manuscripts
    
    # Run synchronously for now (can be shifted to background TM if it takes > 30s)
    try:
        result = compare_manuscripts(documents)
        return jsonify({"success": True, "data": result})
    except Exception as e:
        logger.exception("compare_manuscripts failed")
        return jsonify({"success": False, "error": str(e)}), 500

@paper_rehab_bp.route("/paper-lab/<upload_id>/illustrations", methods=["POST"])
async def generate_illustrations(upload_id: str):
    """
    POST /paper-lab/<id>/illustrations
    Body: {"intent": "A diagram describing the methodology", "task_type": "diagram"}
    """
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    data = request.json or {}
    intent = data.get("intent", "")
    task_type = data.get("task_type", "diagram")
    
    if not intent:
        return jsonify({"success": False, "error": "Provide a 'intent' for the illustration"}), 400

    content = upload.get("full_text", "")

    from ..services.ais.paperbanana_service import generate_paperbanana_illustration
    
    try:
        result = await generate_paperbanana_illustration(content, intent, task_type)
        return jsonify({"success": True, "data": result})
    except Exception as e:
        logger.exception("paperbanana illustration failed")
        return jsonify({"success": False, "error": str(e)}), 500


@paper_rehab_bp.route("/paper-lab/<upload_id>/literature-review", methods=["POST"])
def grounded_literature_review(upload_id):
    """Return verified literature suggestions tied to a focus area."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    data = request.get_json(silent=True) or {}
    focus = (data.get("focus") or "literature review").strip()
    from ..services.ais.paper_orchestra_service import PaperOrchestraService

    result = PaperOrchestraService().grounded_literature_review(upload, focus)
    _update_upload_metadata(upload_id, lambda meta: meta.update({
        "last_grounded_literature_review": result,
        "grounded_literature_history": [*(meta.get("grounded_literature_history") or []), {
            "focus": focus,
            "ready": result.get("ready", False),
            "suggestion_count": len(result.get("suggestions", [])),
            "verified_count": len([item for item in result.get("suggestions", []) if item.get("verified")]),
            "created_at": datetime.now().isoformat(),
        }][-20:],
    }))
    return jsonify({"success": True, "data": result})


@paper_rehab_bp.route("/paper-lab/<upload_id>/refine-section", methods=["POST"])
def refine_section(upload_id):
    """Return a section-scoped revision package with structured diff metadata."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    data = request.get_json(silent=True) or {}
    action = (data.get("action") or "improve_introduction").strip()
    visualization_plan = data.get("visualization_plan")
    from ..services.ais.paper_orchestra_service import PaperOrchestraService

    result = PaperOrchestraService().refine_section(upload, action, visualization_plan)
    refinement_id = f"refine_{uuid.uuid4().hex[:10]}"
    result["refinement_id"] = refinement_id
    result["applied"] = False
    _update_upload_metadata(upload_id, lambda meta: meta.update({
        "last_section_refinement": result,
        "section_refinement_history": [*(meta.get("section_refinement_history") or []), {
            "refinement_id": refinement_id,
            "action": action,
            "section": result.get("section"),
            "summary": result.get("diff", {}).get("summary", ""),
            "created_at": datetime.now().isoformat(),
            "applied": False,
        }][-20:],
    }))
    return jsonify({"success": True, "data": result})


@paper_rehab_bp.route("/paper-lab/<upload_id>/apply-refinement", methods=["POST"])
def apply_refinement(upload_id):
    """Apply a previously generated refinement to the current manuscript draft."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    data = request.get_json(silent=True) or {}
    refinement = data.get("refinement") or {}
    refinement_id = (refinement.get("refinement_id") or data.get("refinement_id") or "").strip()
    section = (refinement.get("section") or "").strip()
    revised_text = refinement.get("revised_text") or ""
    action = refinement.get("action") or "manual_apply"
    if not refinement_id or not section or not revised_text:
        return jsonify({"success": False, "error": "refinement_id, section, and revised_text are required"}), 400

    updated_upload = _apply_section_revision_to_upload(upload, section, revised_text, action, refinement_id)
    _update_upload_metadata(upload_id, lambda meta: meta.update({
        "last_applied_refinement_id": refinement_id,
        "section_refinement_history": [
            {
                **item,
                "applied": True if item.get("refinement_id") == refinement_id else item.get("applied", False),
                "applied_at": datetime.now().isoformat() if item.get("refinement_id") == refinement_id else item.get("applied_at"),
            }
            for item in (meta.get("section_refinement_history") or [])
        ],
    }))

    return jsonify({
        "success": True,
        "data": {
            "upload_id": upload_id,
            "refinement_id": refinement_id,
            "section": section,
            "current_draft": updated_upload.get("current_draft", ""),
            "sections": updated_upload.get("sections", []),
        },
    })


@paper_rehab_bp.route("/paper-lab/<upload_id>/draft-history", methods=["GET"])
def draft_history(upload_id):
    """Return persisted draft/refinement/literature activity history for an upload."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    metadata = upload.get("metadata") or {}
    return jsonify({
        "success": True,
        "data": {
            "upload_id": upload_id,
            "applied_refinements": list(metadata.get("applied_section_refinements") or []),
            "section_refinement_history": list(metadata.get("section_refinement_history") or []),
            "grounded_literature_history": list(metadata.get("grounded_literature_history") or []),
            "last_applied_refinement_id": metadata.get("last_applied_refinement_id"),
        },
    })


@paper_rehab_bp.route("/paper-lab/<upload_id>/revert-refinement", methods=["POST"])
def revert_refinement(upload_id):
    """Revert a previously applied refinement from manuscript history."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    data = request.get_json(silent=True) or {}
    refinement_id = (data.get("refinement_id") or "").strip()
    if not refinement_id:
        return jsonify({"success": False, "error": "refinement_id is required"}), 400

    updated_upload = _revert_section_refinement(upload, refinement_id)
    if not updated_upload:
        return jsonify({"success": False, "error": "Refinement not found"}), 404

    return jsonify({
        "success": True,
        "data": {
            "upload_id": upload_id,
            "reverted_refinement_id": refinement_id,
            "current_draft": updated_upload.get("current_draft", ""),
            "sections": updated_upload.get("sections", []),
        },
    })


@paper_rehab_bp.route("/paper-lab/<upload_id>/graphical-abstract", methods=["POST"])
def graphical_abstract(upload_id):
    """Generate and persist a graphical abstract artifact."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    data = request.get_json(silent=True) or {}
    layout_mode = (data.get("layout_mode") or "process_summary").strip()
    from ..services.ais.paper_orchestra_service import PaperOrchestraService
    from ..services.ais.visualization_artifact_service import VisualizationArtifactService

    payload = PaperOrchestraService().generate_graphical_abstract(upload, layout_mode)
    artifact = VisualizationArtifactService().create_or_replace_by_title(
        upload_id=upload_id,
        artifact_type="graphical_abstract",
        intent="summarize",
        title=payload.get("title", "Graphical Abstract"),
        payload={
            "rendering": {"engine": "html", "spec": payload.get("html")},
            "layout_mode": payload.get("layout_mode"),
            "export_formats": payload.get("export_formats", ["html", "json"]),
            "assumptions": payload.get("assumptions", []),
        },
        provenance=_artifact_provenance(upload, "graphical_abstract"),
        status="needs_input" if payload.get("assumptions") else "ready",
    )
    return jsonify({"success": True, "data": artifact})


@paper_rehab_bp.route("/paper-lab/<upload_id>/slide-starter", methods=["POST"])
def slide_starter(upload_id):
    """Generate and persist a scientific slide starter artifact."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    from ..services.ais.paper_orchestra_service import PaperOrchestraService
    from ..services.ais.visualization_artifact_service import VisualizationArtifactService

    artifact_service = VisualizationArtifactService()
    artifacts = artifact_service.list_for_upload(upload_id)
    payload = PaperOrchestraService().generate_slide_starter(upload, artifacts)
    artifact = artifact_service.create_or_replace_by_title(
        upload_id=upload_id,
        artifact_type="slide",
        intent="summarize",
        title=payload.get("title", "Slide Starter"),
        payload={
            "slides": payload.get("slides", []),
            "export_formats": payload.get("export_formats", ["json"]),
            "assumptions": [] if artifacts else ["Select visuals to enrich the generated slides."],
        },
        provenance=_artifact_provenance(upload, "slide_starter"),
        status="ready" if artifacts else "needs_input",
    )
    return jsonify({"success": True, "data": artifact})


@paper_rehab_bp.route("/paper-lab/<upload_id>/poster-starter", methods=["POST"])
def poster_starter(upload_id):
    """Generate and persist a poster starter artifact."""
    upload = _load_upload(upload_id)
    if not upload:
        return jsonify({"success": False, "error": "Upload not found"}), 404

    from ..services.ais.paper_orchestra_service import PaperOrchestraService
    from ..services.ais.visualization_artifact_service import VisualizationArtifactService

    artifact_service = VisualizationArtifactService()
    artifacts = artifact_service.list_for_upload(upload_id)
    payload = PaperOrchestraService().generate_poster_starter(upload, artifacts)
    artifact = artifact_service.create_or_replace_by_title(
        upload_id=upload_id,
        artifact_type="poster_panel",
        intent="summarize",
        title=payload.get("title", "Poster Starter"),
        payload={
            "panels": payload.get("panels", []),
            "export_formats": payload.get("export_formats", ["json"]),
            "assumptions": [] if artifacts else ["Choose artifact placements for the poster layout."],
        },
        provenance=_artifact_provenance(upload, "poster_starter"),
        status="ready" if artifacts else "needs_input",
    )
    return jsonify({"success": True, "data": artifact})
