"""
Response to Reviewers Generator
Produces a formal academic rebuttal letter from Paper Rehabilitation review data.
Supports: markdown, .docx, and HTML output.
"""

import json
import logging
import re
import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


def generate_response_to_reviewers(
    paper_title: str,
    review_rounds: List[Dict],
    score_progression: List[float],
    source_audit: Dict = None,
    gap_fill_papers: List[Dict] = None,
    novelty_boost: str = "",
    llm=None,
) -> Dict[str, str]:
    """
    Generate a formal Response to Reviewers document.

    Returns: {
        "markdown": str,   # Full markdown document
        "summary": str,    # One-paragraph summary
        "stats": dict,     # Counts of accepted/rebutted/deferred
    }
    """
    gap_fill_papers = gap_fill_papers or []

    # ── Collect all reviewer points across all rounds ──
    all_points = []
    for rd in review_rounds:
        round_num = rd.get("round_num", "?")
        review = rd.get("review", {})
        revision = rd.get("revision", {})

        weaknesses = review.get("top_weaknesses", review.get("all_weaknesses", []))
        triage = revision.get("triage", [])
        responses = revision.get("response_to_reviewers", [])

        # Build triage lookup
        triage_map = {}
        for t in triage:
            key = t.get("weakness", "")[:60]
            triage_map[key] = t
        for r in responses:
            key = r.get("weakness", "")[:60]
            if key not in triage_map:
                triage_map[key] = r

        for w in weaknesses:
            desc = w.get("description", "")
            key = desc[:60]
            t = triage_map.get(key, {})

            all_points.append({
                "round": round_num,
                "reviewer": w.get("flagged_by", "Reviewer"),
                "section": w.get("section", "general"),
                "severity": w.get("severity", "minor"),
                "description": desc,
                "suggestion": w.get("suggestion", ""),
                "action": t.get("action", "accept"),
                "response": t.get("justification", t.get("response", "")),
            })

    # ── Statistics ──
    total = len(all_points)
    accepted = sum(1 for p in all_points if p["action"] == "accept")
    rebutted = sum(1 for p in all_points if p["action"] == "rebut")
    deferred = sum(1 for p in all_points if p["action"] == "defer")
    initial_score = score_progression[0] if score_progression else 0
    final_score = score_progression[-1] if score_progression else 0

    # ── Build markdown ──
    lines = []
    lines.append(f"# Response to Reviewers")
    lines.append("")
    lines.append(f"**Paper:** {paper_title}")
    lines.append(f"**Date:** {datetime.now().strftime('%B %d, %Y')}")
    lines.append(f"**Review Rounds:** {len(review_rounds)}")
    lines.append(f"**Score Progression:** {' → '.join(str(round(s, 1)) for s in score_progression)}")
    lines.append("")
    lines.append("---")
    lines.append("")

    # Summary
    lines.append("## Summary of Changes")
    lines.append("")
    lines.append(
        f"We thank the reviewers for their thorough and constructive feedback. "
        f"In total, **{total} points** were raised across {len(review_rounds)} review rounds. "
        f"We have **accepted {accepted}** suggestions and made corresponding revisions, "
        f"**rebutted {rebutted}** points with justification, "
        f"and **deferred {deferred}** as beyond the current scope. "
    )
    if initial_score and final_score:
        lines.append(
            f"The overall quality score improved from **{initial_score:.1f}** to **{final_score:.1f}** "
            f"(+{final_score - initial_score:.1f})."
        )
    lines.append("")

    # Source audit summary
    if source_audit:
        verified = len(source_audit.get("verified", []))
        unverified = len(source_audit.get("unverified", []))
        lines.append("### Reference Quality")
        lines.append(
            f"A systematic citation audit was conducted: "
            f"**{verified}** references verified, **{unverified}** unverified. "
        )
        if gap_fill_papers:
            lines.append(
                f"We have added **{len(gap_fill_papers)} new references** "
                f"identified through automated literature search to address reviewer concerns."
            )
        lines.append("")

    # Novelty reframing
    if novelty_boost:
        lines.append("### Novelty Reframing")
        lines.append("")
        # Extract just the first section of the novelty boost
        boost_lines = novelty_boost.strip().split("\n")
        for bl in boost_lines[:10]:
            lines.append(bl)
        lines.append("")

    lines.append("---")
    lines.append("")

    # ── Point-by-point responses ──
    lines.append("## Point-by-Point Response")
    lines.append("")

    # Group by reviewer
    by_reviewer = {}
    for p in all_points:
        by_reviewer.setdefault(p["reviewer"], []).append(p)

    for reviewer, points in by_reviewer.items():
        lines.append(f"### {reviewer}")
        lines.append("")

        for i, p in enumerate(points, 1):
            severity_icon = {"fatal": "🔴", "major": "🟡", "minor": "⚪"}.get(p["severity"], "⚪")
            action_tag = {
                "accept": "**[ACCEPTED]**",
                "rebut": "**[REBUTTED]**",
                "defer": "**[DEFERRED]**",
            }.get(p["action"], "**[NOTED]**")

            lines.append(f"**{i}. {severity_icon} [{p['severity'].upper()}] {p['section'].replace('_', ' ').title()}** (Round {p['round']})")
            lines.append("")
            lines.append(f"> *Reviewer:* {p['description']}")
            if p["suggestion"]:
                lines.append(f">")
                lines.append(f"> *Suggestion:* {p['suggestion']}")
            lines.append("")
            lines.append(f"{action_tag}")
            if p["response"]:
                lines.append(f"  {p['response']}")
            elif p["action"] == "accept":
                lines.append(f"  We have revised the {p['section']} section to address this concern.")
            elif p["action"] == "rebut":
                lines.append(f"  We respectfully disagree with this assessment. The current approach is justified because...")
            lines.append("")

        lines.append("---")
        lines.append("")

    # ── New references added ──
    if gap_fill_papers:
        lines.append("## Newly Added References")
        lines.append("")
        for i, p in enumerate(gap_fill_papers, 1):
            doi_text = f" DOI: {p['doi']}" if p.get("doi") else ""
            lines.append(f"{i}. {p.get('title', 'Untitled')} ({p.get('year', '?')}){doi_text}")
        lines.append("")

    # Footer
    lines.append("---")
    lines.append(f"*Generated by OSSR Paper Rehabilitation Pipeline — {datetime.now().strftime('%Y-%m-%d %H:%M')}*")

    markdown = "\n".join(lines)

    return {
        "markdown": markdown,
        "summary": (
            f"Addressed {total} reviewer points: {accepted} accepted, "
            f"{rebutted} rebutted, {deferred} deferred. "
            f"Score: {initial_score:.1f} → {final_score:.1f}."
        ),
        "stats": {
            "total_points": total,
            "accepted": accepted,
            "rebutted": rebutted,
            "deferred": deferred,
            "initial_score": round(initial_score, 1),
            "final_score": round(final_score, 1),
            "rounds": len(review_rounds),
            "new_references": len(gap_fill_papers),
        },
    }


def response_to_docx(response_md: str, paper_title: str) -> bytes:
    """Convert the Response to Reviewers markdown to a .docx file (returns bytes)."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    lines = response_md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line or line == "---":
            i += 1
            continue

        # Headings
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line.lstrip("# "), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line.lstrip("# "), level=1)
        elif line.startswith("### "):
            doc.add_heading(line.lstrip("# "), level=2)
        # Block quotes (reviewer comments)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(36)
            text = line.lstrip("> ").strip()
            # Handle italic markers
            text = text.replace("*Reviewer:*", "").replace("*Suggestion:*", "Suggestion:")
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # Bold tags [ACCEPTED] etc.
        elif line.startswith("**["):
            p = doc.add_paragraph()
            tag = line.strip("*").strip()
            if "ACCEPTED" in tag:
                run = p.add_run(tag)
                run.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            elif "REBUTTED" in tag:
                run = p.add_run(tag)
                run.bold = True
                run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
            elif "DEFERRED" in tag:
                run = p.add_run(tag)
                run.bold = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            else:
                run = p.add_run(tag)
                run.bold = True
        # Bold numbered points
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        # Regular text
        else:
            clean = line.strip()
            # Strip markdown bold/italic
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            if clean:
                doc.add_paragraph(clean)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
