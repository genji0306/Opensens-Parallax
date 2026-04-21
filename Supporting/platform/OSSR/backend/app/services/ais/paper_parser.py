"""
Agent AiS — Paper Parser

Document-ingestion router for Paper Lab and paper crawling.

Supported input types:
- .pdf
- .doc
- .docx
- .txt
- .md / .markdown

Parsing strategy:
- route by file type
- use specialized adapters where available
- normalize all outputs into a canonical ParsedDocumentV2 schema
- fall back to lightweight local parsing when richer engines are unavailable
"""

from __future__ import annotations

import json
import logging
import re
import shutil
import subprocess
import tempfile
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

from ...models.ais_models import PaperSection

logger = logging.getLogger(__name__)

_VI_CHARS = set(
    "àáảãạăắằẳẵặâấầẩẫậèéẻẽẹêếềểễệìíỉĩịòóỏõọôốồổỗộơớờởỡợ"
    "ùúủũụưứừửữựỳýỷỹỵđ"
    "ÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬÈÉẺẼẸÊẾỀỂỄỆÌÍỈĨỊÒÓỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢ"
    "ÙÚỦŨỤƯỨỪỬỮỰỲÝỶỸỴĐ"
)

_SECTION_ALIASES = {
    "abstract": "abstract",
    "tóm tắt": "abstract",
    "introduction": "introduction",
    "giới thiệu": "introduction",
    "mở đầu": "introduction",
    "đặt vấn đề": "introduction",
    "background": "background",
    "cơ sở lý thuyết": "background",
    "tổng quan": "background",
    "literature review": "related_work",
    "related work": "related_work",
    "tài liệu tham khảo": "references",
    "methodology": "methodology",
    "methods": "methodology",
    "phương pháp": "methodology",
    "phương pháp nghiên cứu": "methodology",
    "results": "results",
    "kết quả": "results",
    "kết quả và thảo luận": "results",
    "discussion": "discussion",
    "thảo luận": "discussion",
    "bàn luận": "discussion",
    "conclusion": "conclusion",
    "kết luận": "conclusion",
    "references": "references",
    "bibliography": "references",
    "tài liệu": "references",
    "acknowledgments": "acknowledgments",
    "acknowledgements": "acknowledgments",
}

_DOI_RE = re.compile(r"10\.\d{4,9}/[^\s,;}\]]+")
_CITE_RE = re.compile(r"\(([A-Z][a-z]+(?:\s+(?:et\s+al\.?|&|and)\s+[A-Z][a-z]+)?),?\s*(\d{4})\)")


@dataclass
class DocumentBlock:
    block_id: str
    type: str
    text: str
    page: int = 1
    bbox: Optional[List[float]] = None
    section: str = ""
    confidence: Optional[float] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "block_id": self.block_id,
            "type": self.type,
            "text": self.text,
            "page": self.page,
            "bbox": self.bbox,
            "section": self.section,
            "confidence": self.confidence,
        }


@dataclass
class DocumentTable:
    table_id: str
    page: int
    title: str = ""
    html: str = ""
    markdown: str = ""
    bbox: Optional[List[float]] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "table_id": self.table_id,
            "page": self.page,
            "title": self.title,
            "html": self.html,
            "markdown": self.markdown,
            "bbox": self.bbox,
        }


@dataclass
class DocumentFigure:
    figure_id: str
    page: int
    caption: str = ""
    title: str = ""
    bbox: Optional[List[float]] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "figure_id": self.figure_id,
            "page": self.page,
            "caption": self.caption,
            "title": self.title,
            "bbox": self.bbox,
        }


@dataclass
class DocumentFormula:
    formula_id: str
    page: int
    latex: str = ""
    bbox: Optional[List[float]] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "formula_id": self.formula_id,
            "page": self.page,
            "latex": self.latex,
            "bbox": self.bbox,
        }


@dataclass
class ParsedDocumentV2:
    document_id: str
    source_path: str
    source_filename: str
    source_type: str
    mime_type: str
    parser_engine: str
    parser_mode: str
    ocr_used: bool
    language: str
    title: str
    authors: List[str] = field(default_factory=list)
    abstract: str = ""
    pages: List[Dict[str, Any]] = field(default_factory=list)
    sections: List[PaperSection] = field(default_factory=list)
    blocks: List[DocumentBlock] = field(default_factory=list)
    tables: List[DocumentTable] = field(default_factory=list)
    figures: List[DocumentFigure] = field(default_factory=list)
    formulas: List[DocumentFormula] = field(default_factory=list)
    citations: List[Dict[str, Any]] = field(default_factory=list)
    markdown: str = ""
    plain_text: str = ""
    bbox_index: Dict[str, Any] = field(default_factory=dict)
    parse_warnings: List[str] = field(default_factory=list)
    quality_scores: Dict[str, Any] = field(default_factory=dict)
    enrichment: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "document_id": self.document_id,
            "source_path": self.source_path,
            "source_filename": self.source_filename,
            "source_type": self.source_type,
            "mime_type": self.mime_type,
            "parser_engine": self.parser_engine,
            "parser_mode": self.parser_mode,
            "ocr_used": self.ocr_used,
            "language": self.language,
            "title": self.title,
            "authors": self.authors,
            "abstract": self.abstract,
            "pages": self.pages,
            "sections": [s.to_dict() for s in self.sections],
            "blocks": [b.to_dict() for b in self.blocks],
            "tables": [t.to_dict() for t in self.tables],
            "figures": [f.to_dict() for f in self.figures],
            "formulas": [f.to_dict() for f in self.formulas],
            "citations": self.citations,
            "markdown": self.markdown,
            "plain_text": self.plain_text,
            "bbox_index": self.bbox_index,
            "parse_warnings": self.parse_warnings,
            "quality_scores": self.quality_scores,
            "enrichment": self.enrichment,
        }


@dataclass
class ParsedPaper:
    source_path: str
    title: str
    language: str
    detected_field: str
    sections: List[PaperSection] = field(default_factory=list)
    raw_references: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    full_text: str = ""
    parsed_document: Optional[ParsedDocumentV2] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "source_path": self.source_path,
            "title": self.title,
            "language": self.language,
            "detected_field": self.detected_field,
            "sections": [s.to_dict() for s in self.sections],
            "raw_references": self.raw_references,
            "metadata": self.metadata,
            "word_count": self.metadata.get("word_count", 0),
            "section_count": len(self.sections),
            "parsed_document": self.parsed_document.to_dict() if self.parsed_document else None,
        }


class PaperParser:
    """Routed document parser for PDF, DOC, DOCX, and text inputs."""

    FIELD_KEYWORDS: Dict[str, List[str]] = {
        "mechanical engineering": ["stress", "strain", "deformation", "tensile", "compressive", "elastic", "modulus", "fatigue", "fracture", "FEM", "FEA", "finite element", "ứng suất", "biến dạng", "cơ học"],
        "materials science": ["alloy", "composite", "microstructure", "crystalline", "nanostructure", "polymer", "ceramic", "coating", "vật liệu", "hợp kim", "vi cấu trúc"],
        "thermal engineering": ["temperature", "heat transfer", "thermal conductivity", "convection", "radiation", "cooling", "heating", "nhiệt độ", "truyền nhiệt", "dẫn nhiệt", "tản nhiệt"],
        "electrical engineering": ["impedance", "circuit", "voltage", "current", "sensor", "electrode", "capacitance", "resistance", "trở kháng", "điện trở", "cảm biến"],
        "civil engineering": ["concrete", "reinforcement", "beam", "column", "foundation", "structural", "load", "bê tông", "cốt thép", "kết cấu"],
        "physics": ["quantum", "relativity", "photon", "electron", "wave", "particle", "magnetic", "electric field"],
        "chemistry": ["reaction", "catalyst", "molecular", "compound", "synthesis", "oxidation", "reduction", "phản ứng", "xúc tác"],
        "computer science": ["algorithm", "machine learning", "neural network", "deep learning", "classification", "regression", "optimization", "thuật toán", "học máy"],
        "energy systems": ["battery", "discharge", "capacity", "electrolyte", "cycling", "grid", "renewable", "energy storage"],
    }

    def parse(self, file_path: str) -> ParsedPaper:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Paper not found: {file_path}")

        ext = path.suffix.lower()
        if ext == ".pdf":
            parsed_doc = self._parse_pdf(path)
        elif ext == ".docx":
            parsed_doc = self._parse_docx_document(path)
        elif ext == ".doc":
            parsed_doc = self._parse_doc(path)
        elif ext in (".txt", ".md", ".markdown"):
            parsed_doc = self._parse_text_document(path)
        else:
            raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .doc, .docx, .txt, .md")

        parsed_doc.enrichment = self._ade_enrich(parsed_doc)
        detected_field = self._detect_field(parsed_doc.plain_text)
        citations = parsed_doc.citations or self._extract_references_from_text(parsed_doc.plain_text, parsed_doc.sections)
        parsed_doc.citations = citations

        metadata = {
            "title": parsed_doc.title,
            "parser": parsed_doc.parser_engine,
            "parser_engine": parsed_doc.parser_engine,
            "parser_mode": parsed_doc.parser_mode,
            "ocr_used": parsed_doc.ocr_used,
            "word_count": len(parsed_doc.plain_text.split()),
            "section_count": len(parsed_doc.sections),
            "reference_count": len(citations),
            "file_extension": ext,
            "file_size_kb": round(path.stat().st_size / 1024, 1),
            "document_schema": "ParsedDocumentV2",
            "parse_quality": parsed_doc.quality_scores,
            "parse_warnings": parsed_doc.parse_warnings,
            "figures": [figure.to_dict() for figure in parsed_doc.figures],
            "tables": [table.to_dict() for table in parsed_doc.tables],
            "formulas": [formula.to_dict() for formula in parsed_doc.formulas],
            "parsed_document_v2": parsed_doc.to_dict(),
            "enrichment": parsed_doc.enrichment,
        }

        return ParsedPaper(
            source_path=str(path.resolve()),
            title=parsed_doc.title or path.stem,
            language=parsed_doc.language,
            detected_field=detected_field,
            sections=parsed_doc.sections,
            raw_references=citations,
            metadata=metadata,
            full_text=parsed_doc.plain_text,
            parsed_document=parsed_doc,
        )

    def _parse_text_document(self, path: Path) -> ParsedDocumentV2:
        text = path.read_text(encoding="utf-8", errors="replace")
        title, sections = self._sections_from_text(text, path.stem)
        return self._build_document(
            path=path,
            title=title or path.stem,
            sections=sections,
            plain_text="\n\n".join(section.content for section in sections).strip(),
            parser_engine="legacy_text",
            parser_mode="native",
            ocr_used=False,
            pages=[{"page": 1, "text": text[:8000]}],
        )

    def _parse_docx_document(self, path: Path) -> ParsedDocumentV2:
        from docx import Document

        doc = Document(str(path))
        has_heading_styles = any("heading" in (p.style.name or "").lower() for p in doc.paragraphs if p.text.strip())
        title = ""
        current_heading = "body"
        current_content: List[str] = []
        sections: List[PaperSection] = []
        blocks: List[DocumentBlock] = []
        table_index = 1
        tables: List[DocumentTable] = []

        for para_index, para in enumerate(doc.paragraphs, start=1):
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower()
            is_heading = "heading" in style_name
            if not has_heading_styles and not is_heading:
                is_bold = bool(para.runs) and all(run.bold for run in para.runs if run.text.strip())
                is_short = len(text.split()) <= 8
                is_caps = text.isupper() and len(text) > 3
                matches_keyword = text.lower().strip(":. ") in _SECTION_ALIASES
                if (is_bold and is_short) or is_caps or matches_keyword:
                    is_heading = True

            if is_heading:
                if current_content:
                    sections.append(self._make_section(current_heading, current_content))
                    current_content = []
                current_heading = text
                if not title:
                    title = text
                blocks.append(self._make_block("heading", text, 1, current_heading))
            else:
                current_content.append(text)
                blocks.append(self._make_block("paragraph", text, 1, current_heading))

        if current_content:
            sections.append(self._make_section(current_heading, current_content))

        for table in doc.tables:
            rows = []
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(values) + " |")
            md_table = "\n".join(rows)
            tables.append(DocumentTable(table_id=f"table_{table_index}", page=1, markdown=md_table))
            table_index += 1

        if not sections:
            full = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            sections = [PaperSection(name="body", content=full, word_count=len(full.split()))]
            title = title or path.stem

        plain_text = "\n\n".join(section.content for section in sections).strip()
        return self._build_document(
            path=path,
            title=title or path.stem,
            sections=sections,
            plain_text=plain_text,
            parser_engine="python-docx",
            parser_mode="native",
            ocr_used=False,
            blocks=blocks,
            tables=tables,
            pages=[{"page": 1, "text": plain_text[:8000]}],
        )

    def _parse_doc(self, path: Path) -> ParsedDocumentV2:
        converted = self._convert_doc_to_docx(path)
        parsed = self._parse_docx_document(converted)
        parsed.parser_engine = "doc_conversion+python-docx"
        parsed.parser_mode = "converted"
        parsed.parse_warnings.append("Legacy .doc converted before parsing.")
        return parsed

    def _parse_pdf(self, path: Path) -> ParsedDocumentV2:
        engine_order = [
            self._try_opendataloader_pdf,
            self._try_mineru_pdf,
            self._try_pypdf_pdf,
        ]

        best_doc: Optional[ParsedDocumentV2] = None
        for parser_fn in engine_order:
            candidate = parser_fn(path)
            if not candidate:
                continue
            if best_doc is None or self._quality_score_total(candidate) > self._quality_score_total(best_doc):
                best_doc = candidate
            if self._quality_score_total(candidate) >= 0.72:
                break

        if best_doc is None:
            raise ValueError("No PDF parser backend available. Install pypdf, OpenDataLoader PDF, or MinerU.")

        if self._quality_score_total(best_doc) < 0.55:
            best_doc.parse_warnings.append("PDF parse quality is low; OCR or richer parser backend is recommended.")
        return best_doc

    def _try_pypdf_pdf(self, path: Path) -> Optional[ParsedDocumentV2]:
        try:
            from pypdf import PdfReader
        except Exception:
            return None

        try:
            reader = PdfReader(str(path))
            page_texts = []
            blocks: List[DocumentBlock] = []
            for page_index, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                page_texts.append({"page": page_index, "text": text[:12000]})
                for part in [segment.strip() for segment in text.split("\n") if segment.strip()]:
                    blocks.append(self._make_block("paragraph", part, page_index, ""))
            plain_text = "\n\n".join(page.get("text", "") for page in page_texts).strip()
            title, sections = self._sections_from_text(plain_text, path.stem)
            return self._build_document(
                path=path,
                title=title or path.stem,
                sections=sections,
                plain_text=plain_text,
                parser_engine="pypdf",
                parser_mode="digital_pdf",
                ocr_used=False,
                blocks=blocks,
                pages=page_texts,
            )
        except Exception as exc:
            logger.warning("pypdf parse failed for %s: %s", path.name, exc)
            return None

    def _try_opendataloader_pdf(self, path: Path) -> Optional[ParsedDocumentV2]:
        """
        Optional OpenDataLoader PDF adapter.
        Current implementation uses a local command wrapper if available and
        otherwise returns None.
        """
        cmd = shutil.which("opendataloader-pdf")
        if not cmd:
            return None

        with tempfile.TemporaryDirectory() as tmpdir:
            out_path = Path(tmpdir) / "parsed.json"
            try:
                subprocess.run(
                    [cmd, str(path), "--output", str(out_path)],
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=120,
                )
                if not out_path.exists():
                    return None
                data = json.loads(out_path.read_text(encoding="utf-8"))
                markdown = str(data.get("markdown", ""))
                title, sections = self._sections_from_text(markdown, path.stem)
                return self._build_document(
                    path=path,
                    title=title or path.stem,
                    sections=sections,
                    plain_text=markdown,
                    parser_engine="opendataloader_pdf",
                    parser_mode="layout_aware",
                    ocr_used=bool(data.get("ocr_used", False)),
                    pages=data.get("pages", []),
                    tables=[DocumentTable(table_id=f"table_{idx}", page=int(table.get("page", 1)), html=str(table.get("html", "")), markdown=str(table.get("markdown", ""))) for idx, table in enumerate(data.get("tables", []), start=1)],
                    figures=[DocumentFigure(figure_id=f"figure_{idx}", page=int(figure.get("page", 1)), caption=str(figure.get("caption", ""))) for idx, figure in enumerate(data.get("figures", []), start=1)],
                    formulas=[DocumentFormula(formula_id=f"formula_{idx}", page=int(formula.get("page", 1)), latex=str(formula.get("latex", ""))) for idx, formula in enumerate(data.get("formulas", []), start=1)],
                    markdown=markdown,
                )
            except Exception as exc:
                logger.warning("OpenDataLoader PDF parse failed for %s: %s", path.name, exc)
                return None

    def _try_mineru_pdf(self, path: Path) -> Optional[ParsedDocumentV2]:
        cmd = shutil.which("mineru")
        if not cmd:
            return None

        with tempfile.TemporaryDirectory() as tmpdir:
            out_dir = Path(tmpdir) / "mineru_out"
            try:
                subprocess.run(
                    [cmd, "-p", str(path), "-o", str(out_dir)],
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=180,
                )
                md_files = list(out_dir.rglob("*.md"))
                if not md_files:
                    return None
                markdown = md_files[0].read_text(encoding="utf-8", errors="replace")
                title, sections = self._sections_from_text(markdown, path.stem)
                return self._build_document(
                    path=path,
                    title=title or path.stem,
                    sections=sections,
                    plain_text=markdown,
                    parser_engine="mineru",
                    parser_mode="scientific_pdf",
                    ocr_used=True,
                    pages=[],
                    markdown=markdown,
                )
            except Exception as exc:
                logger.warning("MinerU parse failed for %s: %s", path.name, exc)
                return None

    def _convert_doc_to_docx(self, path: Path) -> Path:
        converters = [
            ["libreoffice", "--headless", "--convert-to", "docx", str(path), "--outdir"],
            ["soffice", "--headless", "--convert-to", "docx", str(path), "--outdir"],
            ["textutil", "-convert", "docx", str(path), "-output"],
        ]
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            for prefix in converters:
                if not shutil.which(prefix[0]):
                    continue
                try:
                    if prefix[0] == "textutil":
                        out_file = tmp_path / f"{path.stem}.docx"
                        subprocess.run(prefix + [str(out_file)], check=True, capture_output=True, text=True, timeout=120)
                    else:
                        subprocess.run(prefix + [str(tmp_path)], check=True, capture_output=True, text=True, timeout=120)
                        out_file = tmp_path / f"{path.stem}.docx"
                    if out_file.exists():
                        persisted = path.with_suffix(".converted.docx")
                        shutil.copy2(out_file, persisted)
                        return persisted
                except Exception as exc:
                    logger.warning("DOC conversion failed via %s for %s: %s", prefix[0], path.name, exc)
        raise ValueError("Unable to convert .doc file. Install libreoffice, soffice, or textutil support.")

    def _sections_from_text(self, text: str, title_fallback: str) -> tuple[str, List[PaperSection]]:
        lines = text.splitlines()
        title = ""
        sections: List[PaperSection] = []
        current_heading = "body"
        current_content: List[str] = []

        for index, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue
            is_md_heading = stripped.startswith("#")
            is_caps_heading = stripped.isupper() and 3 < len(stripped) < 80
            matches_keyword = stripped.lower().strip(":. ") in _SECTION_ALIASES
            if index == 0 and not title and len(stripped.split()) > 3:
                title = stripped[:180]
            if is_md_heading or is_caps_heading or matches_keyword:
                if current_content:
                    sections.append(self._make_section(current_heading, current_content))
                    current_content = []
                current_heading = stripped.lstrip("#").strip()
                if not title:
                    title = current_heading
            else:
                current_content.append(stripped)

        if current_content:
            sections.append(self._make_section(current_heading, current_content))

        if not sections:
            sections = [PaperSection(name="body", content=text, word_count=len(text.split()))]

        return title or title_fallback, sections

    def _build_document(
        self,
        *,
        path: Path,
        title: str,
        sections: List[PaperSection],
        plain_text: str,
        parser_engine: str,
        parser_mode: str,
        ocr_used: bool,
        pages: Optional[List[Dict[str, Any]]] = None,
        blocks: Optional[List[DocumentBlock]] = None,
        tables: Optional[List[DocumentTable]] = None,
        figures: Optional[List[DocumentFigure]] = None,
        formulas: Optional[List[DocumentFormula]] = None,
        markdown: Optional[str] = None,
    ) -> ParsedDocumentV2:
        sections = sections or [PaperSection(name="body", content=plain_text, word_count=len(plain_text.split()))]
        markdown_text = markdown or self._sections_to_markdown(title, sections)
        abstract = next((section.content for section in sections if section.name == "abstract"), "")
        blocks = blocks or self._blocks_from_sections(sections)
        quality = self._score_parse_quality(title, sections, plain_text, tables or [], figures or [], formulas or [])
        return ParsedDocumentV2(
            document_id=f"doc_{uuid.uuid4().hex[:12]}",
            source_path=str(path.resolve()),
            source_filename=path.name,
            source_type=path.suffix.lower().lstrip("."),
            mime_type=self._mime_type(path.suffix.lower()),
            parser_engine=parser_engine,
            parser_mode=parser_mode,
            ocr_used=ocr_used,
            language=self._detect_language(plain_text),
            title=title,
            authors=[],
            abstract=abstract[:1200],
            pages=pages or [{"page": 1, "text": plain_text[:8000]}],
            sections=sections,
            blocks=blocks,
            tables=tables or [],
            figures=figures or [],
            formulas=formulas or [],
            citations=self._extract_references_from_text(plain_text, sections),
            markdown=markdown_text,
            plain_text=plain_text,
            bbox_index={},
            parse_warnings=[],
            quality_scores=quality,
        )

    def _make_section(self, heading: str, content_lines: List[str]) -> PaperSection:
        content = "\n".join(line for line in content_lines if line).strip()
        return PaperSection(
            name=self._normalize_section_name(heading),
            content=content,
            word_count=len(content.split()),
        )

    def _make_block(self, block_type: str, text: str, page: int, section_heading: str) -> DocumentBlock:
        return DocumentBlock(
            block_id=f"block_{uuid.uuid4().hex[:10]}",
            type=block_type,
            text=text,
            page=page,
            section=self._normalize_section_name(section_heading) if section_heading else "",
        )

    def _blocks_from_sections(self, sections: List[PaperSection]) -> List[DocumentBlock]:
        blocks: List[DocumentBlock] = []
        for section in sections:
            blocks.append(self._make_block("heading", section.name, 1, section.name))
            for paragraph in [item.strip() for item in section.content.split("\n") if item.strip()]:
                blocks.append(self._make_block("paragraph", paragraph, 1, section.name))
        return blocks

    def _sections_to_markdown(self, title: str, sections: List[PaperSection]) -> str:
        lines = [f"# {title}"] if title else []
        for section in sections:
            lines.append(f"## {section.name.replace('_', ' ').title()}")
            lines.append(section.content)
        return "\n\n".join(line for line in lines if line)

    def _score_parse_quality(
        self,
        title: str,
        sections: List[PaperSection],
        plain_text: str,
        tables: List[DocumentTable],
        figures: List[DocumentFigure],
        formulas: List[DocumentFormula],
    ) -> Dict[str, Any]:
        title_confidence = 0.9 if title and title.lower() != "body" else 0.4
        section_coverage = min(1.0, len(sections) / 6.0)
        reference_section = 1.0 if any(section.name == "references" for section in sections) else 0.4
        structure = 1.0 if len(plain_text.split()) > 150 else 0.45
        table_retention = 1.0 if tables else 0.6
        figure_retention = 1.0 if figures else 0.6
        formula_retention = 1.0 if formulas else 0.7
        overall = round((title_confidence + section_coverage + reference_section + structure + table_retention + figure_retention + formula_retention) / 7.0, 3)
        return {
            "title_confidence": round(title_confidence, 3),
            "section_coverage": round(section_coverage, 3),
            "reference_section": round(reference_section, 3),
            "structure": round(structure, 3),
            "table_retention": round(table_retention, 3),
            "figure_retention": round(figure_retention, 3),
            "formula_retention": round(formula_retention, 3),
            "overall": overall,
        }

    def _quality_score_total(self, document: ParsedDocumentV2) -> float:
        return float((document.quality_scores or {}).get("overall", 0.0))

    def _mime_type(self, suffix: str) -> str:
        return {
            ".pdf": "application/pdf",
            ".doc": "application/msword",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".txt": "text/plain",
            ".md": "text/markdown",
            ".markdown": "text/markdown",
        }.get(suffix, "application/octet-stream")

    def _normalize_section_name(self, heading: str) -> str:
        key = heading.lower().strip()
        if key in _SECTION_ALIASES:
            return _SECTION_ALIASES[key]
        for alias, canon in _SECTION_ALIASES.items():
            if alias in key or key in alias:
                return canon
        slug = re.sub(r"[^a-z0-9]+", "_", key).strip("_")
        return slug or "unnamed"

    def _detect_language(self, text: str) -> str:
        if not text:
            return "en"
        sample = text[:5000]
        vi_count = sum(1 for ch in sample if ch in _VI_CHARS)
        ratio = vi_count / max(len(sample), 1)
        if ratio > 0.02:
            ascii_words = len(re.findall(r"\b[a-zA-Z]{3,}\b", sample))
            total_words = len(sample.split())
            en_ratio = ascii_words / max(total_words, 1)
            if en_ratio > 0.5:
                return "mixed"
            return "vi"
        return "en"

    def _detect_field(self, text: str) -> str:
        text_lower = text.lower()
        scores: Dict[str, int] = {}
        for field_name, keywords in self.FIELD_KEYWORDS.items():
            score = sum(text_lower.count(keyword.lower()) for keyword in keywords)
            if score > 0:
                scores[field_name] = score
        if not scores:
            return "general science"
        return max(scores, key=scores.get)

    def _extract_references_from_text(self, full_text: str, sections: List[PaperSection]) -> List[Dict[str, Any]]:
        refs: List[Dict[str, Any]] = []
        seen_dois = set()
        for match in _DOI_RE.finditer(full_text):
            doi = match.group().rstrip(".")
            if doi not in seen_dois:
                seen_dois.add(doi)
                refs.append({"type": "doi", "doi": doi, "raw": doi})

        ref_section = next((section for section in sections if section.name in {"references", "bibliography"}), None)
        if ref_section:
            for line in ref_section.content.split("\n"):
                line = line.strip()
                if not line or len(line) < 10:
                    continue
                cleaned = re.sub(r"^\[?\d+\]?\.?\s*", "", line).strip()
                doi_match = _DOI_RE.search(cleaned)
                doi = doi_match.group().rstrip(".") if doi_match else ""
                if doi and doi in seen_dois:
                    continue
                if doi:
                    seen_dois.add(doi)
                refs.append({
                    "type": "reference_entry",
                    "doi": doi,
                    "raw": cleaned[:300],
                    "title": self._guess_title_from_ref(cleaned),
                })

        for match in _CITE_RE.finditer(full_text):
            author, year = match.group(1), match.group(2)
            refs.append({
                "type": "author_year",
                "doi": "",
                "raw": f"{author}, {year}",
                "author": author,
                "year": int(year),
            })
        return refs

    def _ade_enrich(self, document: ParsedDocumentV2) -> Dict[str, Any]:
        """
        Optional ADE-style enrichment hook.
        This remains safe when ADE is not installed.
        """
        if document.source_type not in {"pdf", "doc", "docx"}:
            return {"engine": "none", "applied": False}

        try:
            import ade  # type: ignore  # noqa: F401
            return {
                "engine": "ade_python",
                "applied": True,
                "extracted_fields": {
                    "title": document.title,
                    "abstract": document.abstract[:500],
                    "section_names": [section.name for section in document.sections],
                },
            }
        except Exception:
            return {
                "engine": "none",
                "applied": False,
                "fallback_reason": "ADE not installed or configured",
            }

    @staticmethod
    def _guess_title_from_ref(ref_line: str) -> str:
        quoted = re.search(r'["""](.+?)["""]', ref_line)
        if quoted:
            return quoted.group(1)
        parts = ref_line.split(".")
        if len(parts) >= 3:
            candidate = parts[1].strip()
            if 10 < len(candidate) < 200:
                return candidate
        return ""
